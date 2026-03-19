import os
import copy
import streamlit as st
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from datetime import date
import io
import requests

st.set_page_config(page_title="Fattura Generator", layout="wide")

def fmt_price(n):
    """Format number as European: 2.470,– for whole numbers, 0,15 for decimals, -832,85 for negatives."""
    sign = "-" if n < 0 else ""
    abs_n = abs(n)
    formatted = f"{abs_n:,.2f}"
    formatted = formatted.replace(",", "X").replace(".", ",").replace("X", ".")
    if formatted.endswith(",00"):
        # Whole number — add the dash
        formatted = formatted[:-3] + ",–"
    # else: has real decimal digits — no dash
    return f"{sign}{formatted}"

def fmt_qty(n):
    """Format quantity with comma: 1,0"""
    return f"{n:.1f}".replace(".", ",")

# ─────────────────────────────────────────────
# PASSWORD GATE
# ─────────────────────────────────────────────
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if not st.session_state.authenticated:
    st.title("🔒 Fattura Generator")
    pwd = st.text_input("Enter passcode to continue:", type="password")
    if st.button("Login"):
        if pwd == "RAINYEAR":
            st.session_state.authenticated = True
            st.rerun()
        else:
            st.error("❌ Wrong passcode.")
    st.stop()

# ─────────────────────────────────────────────
# SUPABASE
# ─────────────────────────────────────────────
SUPABASE_URL = "https://lztrggttkgvgjouofibd.supabase.co"
SUPABASE_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6Imx6dHJnZ3R0a2d2Z2pvdW9maWJkIiwicm9sZSI6ImFub24iLCJpYXQiOjE3NzMyNDAwNzEsImV4cCI6MjA4ODgxNjA3MX0.tbHCQtGW21C2fXCEu2FGwlsXn4kGUWOGoOqjuYyiC7A"
HEADERS = {
    "apikey": SUPABASE_KEY,
    "Authorization": f"Bearer {SUPABASE_KEY}",
    "Content-Type": "application/json"
}

def get_next_invoice_number(fattura_type):
    year_2digit = date.today().strftime('%y')
    prefix = "INI" if fattura_type == "Fattura Italia" else "INE"
    response = requests.get(f"{SUPABASE_URL}/rest/v1/fatture", headers=HEADERS,
                            params={"select": "invoice_number"})
    try:
        existing = response.json()
        if isinstance(existing, list):
            this_year = [r for r in existing
                         if str(r.get("invoice_number","")).startswith(prefix)
                         and str(r.get("invoice_number","")).endswith(f"/{year_2digit}")]
            next_num = len(this_year) + 1
        else:
            next_num = 1
    except:
        next_num = 1
    return f"{prefix}{next_num:03d}/{year_2digit}"

def save_fattura(invoice_number, client_company, total_amount, currency,
                 address="", zip_code="", city="", region="", country="",
                 date_of_reference=None, note=None, payment_terms=None):
    r = requests.post(
        f"{SUPABASE_URL}/rest/v1/fatture",
        headers={**HEADERS, "Prefer": "return=representation"},
        json={
            "invoice_number": invoice_number,
            "client_company": client_company,
            "total_amount": total_amount,
            "currency": currency,
            "status": "not_sent",
            "address": address, "zip": zip_code, "city": city,
            "region": region, "country": country,
            "date_of_reference": date_of_reference,
            "note": note, "payment_terms": payment_terms,
        }
    )
    if not r.ok:
        st.warning(f"⚠️ Could not save fattura: {r.status_code} {r.text}")
        return None
    data = r.json()
    if isinstance(data, list) and data:
        return data[0].get("id")
    return None

def save_fattura_items(fattura_id, items):
    if not fattura_id: return
    rows = []
    for it in items:
        if not it.get("description","").strip(): continue
        p = PRODUCT_MAP.get(it.get("product_idx", 0), {})
        rows.append({
            "fattura_id":    fattura_id,
            "description":   it.get("description",""),
            "description_it":it.get("description_it",""),
            "qty":           it.get("qty", 0),
            "unit_price":    it.get("unit_price", 0),
            "currency":      it.get("currency","EUR"),
            "net_weight_kg": float(p.get("net_weight_kg") or 0) if p else 0,
            "dimensions":    p.get("dimensions") or "" if p else "",
        })
    if rows:
        requests.post(f"{SUPABASE_URL}/rest/v1/fattura_items",
                      headers={**HEADERS, "Prefer": "return=minimal"}, json=rows)

def load_products():
    try:
        data = requests.get(f"{SUPABASE_URL}/rest/v1/products", headers=HEADERS,
            params={"select":"id,description,description_eng,unit_price_client,unit_price_reseller,category,net_weight_kg,dimensions",
                    "order":"category.asc,created_at.asc"}).json()
        return data if isinstance(data, list) else []
    except: return []

def load_customers():
    try:
        data = requests.get(f"{SUPABASE_URL}/rest/v1/customers", headers=HEADERS,
            params={"select":"id,company_name,contact_name,salutation,address,city,zip,state,country,vat_number",
                    "order":"company_name.asc"}).json()
        return data if isinstance(data, list) else []
    except: return []

def load_delivery_addresses():
    try:
        data = requests.get(f"{SUPABASE_URL}/rest/v1/delivery_addresses", headers=HEADERS,
            params={"select":"*","order":"company_name.asc"}).json()
        return data if isinstance(data, list) else []
    except: return []

def load_delivery_terms():
    try:
        data = requests.get(f"{SUPABASE_URL}/rest/v1/delivery_terms", headers=HEADERS,
            params={"select":"term","order":"created_at.asc"}).json()
        return [r["term"] for r in data] if isinstance(data, list) else []
    except: return []

def load_payment_terms():
    try:
        data = requests.get(f"{SUPABASE_URL}/rest/v1/payment_terms", headers=HEADERS,
            params={"select":"term","order":"created_at.asc"}).json()
        return [r["term"] for r in data] if isinstance(data, list) else []
    except: return []

def save_payment_term(term):
    existing = load_payment_terms()
    if term in existing: return
    requests.post(f"{SUPABASE_URL}/rest/v1/payment_terms", headers=HEADERS, json={"term": term})

def save_delivery_address(company_name, street_address, zip_code, city, country):
    check = requests.get(f"{SUPABASE_URL}/rest/v1/delivery_addresses", headers=HEADERS,
        params={"company_name":f"eq.{company_name}","select":"id"})
    try:
        if isinstance(check.json(), list) and len(check.json()) > 0: return
    except: pass
    requests.post(f"{SUPABASE_URL}/rest/v1/delivery_addresses", headers=HEADERS,
        json={"company_name":company_name,"street_address":street_address,
              "zip_code":zip_code,"city":city,"country":country})

def load_vat_exemptions():
    try:
        data = requests.get(f"{SUPABASE_URL}/rest/v1/vat_exemptions", headers=HEADERS,
            params={"select":"code","order":"created_at.asc"}).json()
        return [r["code"] for r in data if r.get("code")] if isinstance(data, list) else []
    except: return []

def save_vat_exemption(text):
    check = requests.get(f"{SUPABASE_URL}/rest/v1/vat_exemptions", headers=HEADERS,
        params={"code":f"eq.{text}","select":"id"})
    try:
        if isinstance(check.json(), list) and len(check.json()) > 0: return
    except: pass
    requests.post(f"{SUPABASE_URL}/rest/v1/vat_exemptions", headers=HEADERS, json={"code": text})

# NEW: load fatture with status "Fattura di anticipo"
def load_fatture_anticipo():
    try:
        data = requests.get(f"{SUPABASE_URL}/rest/v1/fatture", headers=HEADERS,
            params={"select":"id,invoice_number,client_company,total_amount,currency,date_of_reference",
                    "status":"eq.Fattura di anticipo",
                    "order":"created_at.desc"}).json()
        return data if isinstance(data, list) else []
    except: return []

# ─────────────────────────────────────────────
# DOCX HELPERS
# ─────────────────────────────────────────────
def set_cell_text(cell, text, bold=False, italic=False, font_name="Verdana", font_size=10):
    tc = cell._tc
    paras = tc.findall(qn('w:p'))
    for extra_p in paras[1:]:
        tc.remove(extra_p)
    first_p = cell.paragraphs[0]
    for run in first_p.runs:
        run.text = ""
        rPr = run._r.find(qn('w:rPr'))
        if rPr is not None: run._r.remove(rPr)
    lines = text.split("\n")
    run = first_p.add_run(lines[0])
    run.bold = bold; run.italic = italic
    run.font.name = font_name; run.font.size = Pt(font_size)
    for line in lines[1:]:
        br = OxmlElement("w:br")
        run._r.addnext(br)
        run2 = first_p.add_run(line)
        run2.bold = bold; run2.italic = italic
        run2.font.name = font_name; run2.font.size = Pt(font_size)
        run = run2

def replace_in_table_cell(cell, replacements):
    for para in cell.paragraphs:
        full_text = "".join(run.text for run in para.runs)
        changed = False
        for key, val in replacements.items():
            if key in full_text:
                full_text = full_text.replace(key, val)
                changed = True
        if changed and para.runs:
            para.runs[0].text = full_text
            for run in para.runs[1:]: run.text = ""

def replace_in_paragraph(para, replacements):
    full_text = "".join(run.text for run in para.runs)
    changed = False
    for key, val in replacements.items():
        if key in full_text:
            full_text = full_text.replace(key, val)
            changed = True
    if changed and para.runs:
        para.runs[0].text = full_text
        for run in para.runs[1:]: run.text = ""

def delete_para(para):
    p = para._p
    p.getparent().remove(p)

# ─────────────────────────────────────────────
# LOAD DATA
# ─────────────────────────────────────────────
if "products_db" not in st.session_state:
    st.session_state.products_db = load_products()
if "customers_db" not in st.session_state:
    st.session_state.customers_db = load_customers()
if "delivery_db" not in st.session_state:
    st.session_state.delivery_db = load_delivery_addresses()
if "delivery_terms_db" not in st.session_state:
    st.session_state.delivery_terms_db = load_delivery_terms()
if "payment_terms_db" not in st.session_state:
    st.session_state.payment_terms_db = load_payment_terms()
if "vat_exemptions_db" not in st.session_state:
    st.session_state.vat_exemptions_db = load_vat_exemptions()
if "fatture_anticipo_db" not in st.session_state:
    st.session_state.fatture_anticipo_db = load_fatture_anticipo()

PRODUCTS = st.session_state.products_db
CATEGORIES = []
seen_cats = []
for p in PRODUCTS:
    cat = p.get("category") or "Other"
    if cat not in seen_cats:
        seen_cats.append(cat)
        CATEGORIES.append(cat)

PRODUCT_NAMES = ["— custom item —"]
PRODUCT_MAP   = {}
for cat in CATEGORIES:
    cat_products = [p for p in PRODUCTS if (p.get("category") or "Other") == cat]
    for p in cat_products:
        eng = p.get("description_eng") or p["description"]
        label = eng[:60] + ("…" if len(eng) > 60 else "")
        PRODUCT_MAP[len(PRODUCT_NAMES)] = p
        PRODUCT_NAMES.append(label)

# ─────────────────────────────────────────────
# OPTIONS
# ─────────────────────────────────────────────
CURRENCIES = ["EUR", "USD", "GBP", "CHF", "CNY", "RUB", "— custom —"]
HS_CODES   = ["8453.9000","8453.1000","8466.9195","8464.2019","8451.9000","8451.8030","— custom —"]

# ─────────────────────────────────────────────
# SESSION STATE
# ─────────────────────────────────────────────
if "fattura_line_items" not in st.session_state:
    st.session_state.fattura_line_items = [
        {"product_idx":0,"description":"","description_it":"","details":"",
         "qty":1.0,"unit_price":0.0,"price_type":"Cliente",
         "is_discount":False,"discount_value":0.0,"linked_anticipo":None}
    ]

def add_line():
    st.session_state.fattura_line_items.append(
        {"product_idx":0,"description":"","description_it":"","details":"",
         "qty":1.0,"unit_price":0.0,"price_type":"Cliente",
         "is_discount":False,"discount_value":0.0,"linked_anticipo":None}
    )

def add_discount_line():
    st.session_state.fattura_line_items.append(
        {"product_idx":-1,"description":"DEDUCTION DOWN PAYMENT BY T/T",
         "description_it":"Deduzione per anticipo a mezzo bonifico bancario",
         "details":"","qty":1.0,"unit_price":0.0,"price_type":"Cliente",
         "is_discount":True,"discount_value":0.0,"linked_anticipo":None}
    )

# ─────────────────────────────────────────────
# UI
# ─────────────────────────────────────────────
st.title("🧾 Fattura Generator")

# ── 1. DATE, TYPE & NUMBER ──
st.subheader("1. Date, Type & Invoice Number")
col_d1, col_d2, col_d3 = st.columns(3)
with col_d1:
    selected_date = st.date_input("Date", value=date.today(), format="DD/MM/YYYY")
with col_d2:
    fattura_type = st.radio("Fattura Type", ["Fattura Estero","Fattura Italia"], horizontal=True, key="fattura_type")
with col_d3:
    invoice_number = get_next_invoice_number(fattura_type)
    st.metric("Invoice Number", invoice_number)

formatted_date = selected_date.strftime('%d/%m/%Y')

# ── NOTE ──
note = st.text_input("📝 Note (optional — shown in the app)", placeholder="e.g. Spare parts order, urgent delivery")

# ── 2. CLIENT ──
st.subheader("2. Client")
customers = st.session_state.customers_db
customer_names = ["— new customer —"] + [
    f"{c.get('company_name','')} ({c.get('contact_name','')})" for c in customers
]
col_cust, col_refresh = st.columns([5,1])
with col_cust:
    selected_customer_idx = st.selectbox(
        "Pick existing customer or fill in manually",
        range(len(customer_names)), format_func=lambda x: customer_names[x], key="cust_picker")
with col_refresh:
    st.write("")
    if st.button("🔄", help="Reload customers"):
        st.session_state.customers_db = load_customers(); st.rerun()

if selected_customer_idx > 0:
    cust = customers[selected_customer_idx-1]
    default_company    = cust.get("company_name","")
    default_address    = cust.get("address","")
    default_zip        = cust.get("zip","")
    default_city       = cust.get("city","")
    default_region     = cust.get("state","") or ""
    default_country    = cust.get("country","")
    default_vat        = cust.get("vat_number","")
    default_salutation = cust.get("salutation","Mr.") or "Mr."
    default_full_name  = cust.get("contact_name","") or ""
else:
    default_company=default_address=default_zip=""
    default_city=default_country=default_vat=""; default_region=""
    default_salutation="Mr."; default_full_name=""

include_attn = st.checkbox("Include 'To the attn. of' line?", value=False)
salutation=""; full_name=""
if include_attn:
    col_s, col_n = st.columns([1,3])
    with col_s:
        sal_opts = ["Mr.","Ms.","Dr.","Messrs."]
        sal_idx  = sal_opts.index(default_salutation) if default_salutation in sal_opts else 0
        salutation = st.selectbox("Salutation", sal_opts, index=sal_idx)
    with col_n:
        full_name = st.text_input("Full Name (optional)", value=default_full_name)

company    = st.text_input("Company Name *", value=default_company)
address    = st.text_input("Address", value=default_address)
col3,col4,col5 = st.columns(3)
with col3: zip_code = st.text_input("Zip", value=default_zip)
with col4: city = st.text_input("City", value=default_city)
with col5: region = st.text_input("Region", value=default_region, placeholder="(optional)")
country    = st.text_input("Country", value=default_country)
vat_number = st.text_input("Tax ID / VAT code / Partita IVA", value=default_vat)

# ── 3. DELIVERY ADDRESS ──
st.subheader("3. Delivery Address")
delivery_addresses = st.session_state.delivery_db
delivery_names = ["— same as billing —","— new address —"] + [
    f"{d.get('company_name','')} — {d.get('city','')}" for d in delivery_addresses
]
col_del, col_del_refresh = st.columns([5,1])
with col_del:
    selected_delivery_idx = st.selectbox("Select delivery address",
        range(len(delivery_names)), format_func=lambda x: delivery_names[x], key="delivery_picker")
with col_del_refresh:
    st.write("")
    if st.button("🔄", key="reload_delivery", help="Reload delivery addresses"):
        st.session_state.delivery_db = load_delivery_addresses(); st.rerun()

if selected_delivery_idx == 0:
    del_company=company; del_address=address; del_zip=zip_code
    del_city=city; del_region=region; del_country=country
    st.caption("📦 Delivery address same as billing address")
elif selected_delivery_idx == 1:
    del_company = st.text_input("Delivery Company Name")
    del_address = st.text_input("Delivery Street Address")
    col_dz, col_dc = st.columns(2)
    with col_dz: del_zip = st.text_input("Delivery ZIP")
    with col_dc: del_city = st.text_input("Delivery City")
    del_region  = st.text_input("Delivery Region", placeholder="(optional)")
    del_country = st.text_input("Delivery Country")
    if del_company and st.button("💾 Save this delivery address"):
        save_delivery_address(del_company, del_address, del_zip, del_city, del_country)
        st.session_state.delivery_db = load_delivery_addresses()
        st.success(f"✅ '{del_company}' saved!"); st.rerun()
else:
    d = delivery_addresses[selected_delivery_idx-2]
    del_company=d.get("company_name",""); del_address=d.get("street_address","")
    del_zip=d.get("zip_code",""); del_city=d.get("city","")
    del_region=""; del_country=d.get("country","")
    st.caption(f"📦 {del_company} — {del_address}, {del_zip} {del_city}, {del_country}")

# ── 4. TERMS ──
st.subheader("4. Terms")
col_t1, col_t2 = st.columns(2)
with col_t1:
    delivery_terms_options = st.session_state.delivery_terms_db
    delivery_terms = st.selectbox("Delivery Terms", delivery_terms_options+["— custom —"])
    if delivery_terms == "— custom —":
        delivery_terms = st.text_input("Custom Delivery Terms", placeholder="e.g. DAP Tokyo")
    payment_terms_options = st.session_state.payment_terms_db
    payment = st.selectbox("Payment Terms", payment_terms_options+["— custom —"])
    if payment == "— custom —":
        payment = st.text_input("Custom Payment Terms", placeholder="e.g. 60 days from invoice")
        if payment and payment not in payment_terms_options:
            if st.button("💾 Save this payment term", key="save_pt"):
                save_payment_term(payment)
                st.session_state.payment_terms_db = load_payment_terms()
                st.success(f"✅ '{payment}' saved!"); st.rerun()
with col_t2:
    hs_code = st.selectbox("HS Code", HS_CODES)
    if hs_code == "— custom —":
        hs_code = st.text_input("Custom HS Code")
    vat_options_dynamic = ["— none —"] + st.session_state.vat_exemptions_db + ["— custom —"]
    vat_exemption_choice = st.selectbox("VAT Exemption", vat_options_dynamic)
    if vat_exemption_choice == "— custom —":
        vat_exemption = st.text_input("Custom VAT exemption text", placeholder="e.g. Art. 8 DPR 633/72")
        if vat_exemption and st.button("💾 Save this VAT exemption", key="save_vat"):
            save_vat_exemption(vat_exemption)
            st.session_state.vat_exemptions_db = load_vat_exemptions()
            st.success("✅ Saved!"); st.rerun()
    elif vat_exemption_choice == "— none —":
        vat_exemption = ""
    else:
        vat_exemption = vat_exemption_choice

# ── 5. CURRENCY & PRICE TYPE ──
st.subheader("5. Currency & Price Type")
col_cur, col_pt = st.columns(2)
with col_cur:
    currency_choice = st.selectbox("Currency (ISO)", CURRENCIES)
    currency = st.text_input("ISO currency code", placeholder="e.g. AED") if currency_choice == "— custom —" else currency_choice
with col_pt:
    global_price_type = st.radio("Price type", ["Cliente","Rivenditore"], horizontal=True, key="fattura_price_type")
    if st.session_state.get("_fattura_last_price_type") != global_price_type:
        st.session_state["_fattura_last_price_type"] = global_price_type
        for item in st.session_state.fattura_line_items:
            if item.get("is_discount"): continue
            item["price_type"] = global_price_type
            if item.get("product_idx",0) > 0 and item.get("product_idx") in PRODUCT_MAP:
                pc = item.get("price_client",0.0)
                pr = item.get("price_reseller",0.0)
                item["unit_price"] = pc if global_price_type == "Cliente" else pr
        st.rerun()

# ── 6. LINE ITEMS ──
st.subheader("6. Line Items")
st.caption("Select from catalogue or choose '— custom item —' to type manually.")

# Extra catalogue items (not in products DB, built-in)
EXTRA_ITEMS = [
    # (label_en, label_it, description_en, description_it)
    ("Packing charges",
     "Spese di imballaggio",
     "Packing charges",
     "Spese di imballaggio"),
    ("Packing and shipping charges",
     "Spese di imballaggio e spedizione",
     "Packing and shipping charges",
     "Spese di imballaggio e spedizione"),
    ("Shipping charges",
     "Spese di spedizione",
     "Shipping charges",
     "Spese di spedizione"),
]
EXTRA_ITEM_LABELS = [f"🇬🇧 {e[0]} / 🇮🇹 {e[1]}" for e in EXTRA_ITEMS]
EXTRA_ITEM_OFFSET = len(PRODUCT_NAMES)  # indices >= this are extra items

ALL_ITEM_NAMES = PRODUCT_NAMES + EXTRA_ITEM_LABELS

items_to_remove = []
needs_rerun = False

for i, item in enumerate(st.session_state.fattura_line_items):

    with st.container():

        # ── DISCOUNT LINE ──────────────────────────────────────────────
        if item.get("is_discount"):
            st.markdown(f"**📉 Deduction — Down Payment #{i+1}**")
            st.caption("This deduction reduces the invoice total by the advance amount previously paid by the customer.")

            fatture_anticipo = st.session_state.fatture_anticipo_db
            anticipo_options = ["— select the advance payment invoice —"] + [
                f"{f.get('invoice_number','')} — {f.get('client_company','')} ({f.get('currency','EUR')} {f.get('total_amount',0):,.2f})"
                for f in fatture_anticipo
            ]

            col_a1, col_a2 = st.columns([3, 0.4])
            with col_a1:
                sel_anticipo_idx = st.selectbox(
                    "Reference invoice (advance payment previously received) *",
                    range(len(anticipo_options)),
                    format_func=lambda x: anticipo_options[x],
                    key=f"anticipo_{i}"
                )
                if sel_anticipo_idx > 0:
                    fat = fatture_anticipo[sel_anticipo_idx - 1]
                    inv_num  = fat.get("invoice_number","")
                    inv_date = fat.get("date_of_reference","") or ""
                    # Format date dd/mm/yyyy if ISO
                    if inv_date and "-" in inv_date:
                        parts = inv_date.split("-")
                        if len(parts) == 3:
                            inv_date = f"{parts[2]}/{parts[1]}/{parts[0]}"
                    # EN description
                    item["description"]    = "DEDUCTION DOWN PAYMENT BY T/T"
                    item["details"]        = f"Our invoice no. {inv_num} dtd {inv_date}"
                    # IT description
                    item["description_it"] = "Deduzione per anticipo a mezzo bonifico bancario"
                    item["details_it"]     = f"Nostra fattura no. {inv_num} del {inv_date}"
                    item["linked_anticipo"] = inv_num
                    st.caption(f"🇬🇧 Description: {item['description']} | {item['details']}")
                    st.caption(f"🇮🇹 Descrizione: {item['description_it']} | {item['details_it']}")
                else:
                    item["linked_anticipo"] = None

                # Discount value — must be ≤ 0
                discount_val = st.number_input(
                    f"Discount Value ({currency}) — must be 0 or negative",
                    max_value=0.0,
                    value=float(item.get("discount_value", 0.0)),
                    step=1.0,
                    format="%.2f",
                    key=f"disc_val_{i}"
                )
                item["discount_value"] = discount_val
                item["unit_price"]     = discount_val   # used in grand total
                item["qty"]            = 1.0
                if discount_val != 0.0:
                    st.caption(f"Discount total: {currency} {fmt_price(discount_val)}")

            with col_a2:
                st.write(""); st.write("")
                if st.button("🗑", key=f"fattura_del_{i}"):
                    items_to_remove.append(i)

            st.divider()
            continue

        # ── NORMAL LINE ────────────────────────────────────────────────
        c1, c2, c3, c4 = st.columns([3, 1.5, 1.5, 0.4])
        with c1:
            prod_idx = st.selectbox(
                f"Product #{i+1}",
                range(len(ALL_ITEM_NAMES)),
                format_func=lambda x: ALL_ITEM_NAMES[x],
                key=f"fattura_prod_{i}",
                index=item.get("product_idx", 0)
            )
            if prod_idx != item.get("product_idx"):
                item["product_idx"] = prod_idx
                if 0 < prod_idx < EXTRA_ITEM_OFFSET and prod_idx in PRODUCT_MAP:
                    p = PRODUCT_MAP[prod_idx]
                    item["description"]    = p.get("description_eng") or p["description"]
                    item["description_it"] = p.get("description","")
                    item["price_client"]   = float(p.get("unit_price_client") or 0)
                    item["price_reseller"] = float(p.get("unit_price_reseller") or 0)
                    item["unit_price"] = item["price_client"] if global_price_type == "Cliente" else item["price_reseller"]
                    item["is_extra"] = False
                elif prod_idx >= EXTRA_ITEM_OFFSET:
                    extra = EXTRA_ITEMS[prod_idx - EXTRA_ITEM_OFFSET]
                    item["description"]    = extra[0]
                    item["description_it"] = extra[1]
                    item["price_client"] = item["price_reseller"] = 0.0
                    item["unit_price"] = 0.0
                    item["is_extra"] = True
                else:
                    item["description"]=""; item["description_it"]=""
                    item["unit_price"]=item["price_client"]=item["price_reseller"]=0.0
                    item["is_extra"]=False
                # Delete the number_input widget key so it reinitialises with the new price
                st.session_state.pop(f"fattura_up_{i}", None)
                needs_rerun = True

            # Show caption for catalogue items
            if 0 < prod_idx < EXTRA_ITEM_OFFSET and prod_idx in PRODUCT_MAP:
                it_name = PRODUCT_MAP[prod_idx].get("description","")
                if it_name: st.caption(f"🇮🇹 {it_name}")

            # Show extra item caption
            if prod_idx >= EXTRA_ITEM_OFFSET:
                extra = EXTRA_ITEMS[prod_idx - EXTRA_ITEM_OFFSET]
                st.caption(f"🇬🇧 {extra[0]} / 🇮🇹 {extra[1]}")

            # Custom item fields — fattura app is English-only for the document
            # but stores both EN and IT descriptions for bilingual templates
            if prod_idx == 0:
                item["description"] = st.text_input(
                    "Custom Product Name (EN — shown in document)",
                    value=item.get("description", ""),
                    key=f"fattura_desc_{i}")
                item["description_it"] = st.text_input(
                    "Custom Product Name (IT — for reference)",
                    value=item.get("description_it", ""),
                    key=f"fattura_desc_it_{i}")

            item["details"] = st.text_input(
                "Description / Specs (optional)", value=item.get("details",""), key=f"fattura_details_{i}")

            # Unit price — always editable; pre-filled from DB list price for catalogue items.
            # Key includes prod_idx so Streamlit creates a fresh widget when product changes.
            is_catalogue = 0 < prod_idx < EXTRA_ITEM_OFFSET and prod_idx in PRODUCT_MAP
            db_price = 0.0
            if is_catalogue:
                db_price = item.get("price_client", 0.0) if global_price_type == "Cliente" else item.get("price_reseller", 0.0)

            item["unit_price"] = st.number_input(
                f"Unit Price ({currency})",
                min_value=0.0,
                value=float(item.get("unit_price", db_price if is_catalogue else 0.0)),
                step=0.01, format="%.2f",
                key=f"fattura_up_{i}_{prod_idx}"
            )

            # Discount / surcharge indicator — only for catalogue items with known DB price
            if is_catalogue and db_price > 0:
                entered = float(item.get("unit_price", 0.0))
                if abs(entered - db_price) > 0.001:
                    diff_pct = ((entered - db_price) / db_price) * 100
                    diff_abs = entered - db_price
                    if entered < db_price:
                        st.caption(
                            f"🔴 Discount: −{currency} {fmt_price(abs(diff_abs))} "
                            f"({abs(diff_pct):.1f}% below list price of {currency} {fmt_price(db_price)})"
                        )
                    else:
                        st.caption(
                            f"🟢 Surcharge: +{currency} {fmt_price(diff_abs)} "
                            f"({diff_pct:.1f}% above list price of {currency} {fmt_price(db_price)})"
                        )

        with c2:
            item["qty"] = st.number_input("Qty", min_value=0.0, value=float(item.get("qty",1.0)),
                step=1.0, format="%.1f", key=f"fattura_qty_{i}")
        with c3:
            line_total = item["qty"] * item["unit_price"]
            st.write(f"**Line Total ({currency})**")
            st.write(fmt_price(line_total))
        with c4:
            st.write(""); st.write("")
            if st.button("🗑", key=f"fattura_del_{i}"):
                items_to_remove.append(i)

        line_total = item["qty"] * item["unit_price"]
        st.caption(f"Line total: {currency} {fmt_price(line_total)}")
        st.divider()

for i in sorted(items_to_remove, reverse=True):
    st.session_state.fattura_line_items.pop(i)
if items_to_remove or needs_rerun:
    st.rerun()

# Add line buttons
col_btn1, col_btn2 = st.columns(2)
with col_btn1:
    st.button("➕ Add Line Item", on_click=add_line)
with col_btn2:
    if st.button("➕ Add Discount / Sconto"):
        # Reload anticipo fatture fresh
        st.session_state.fatture_anticipo_db = load_fatture_anticipo()
        add_discount_line()
        st.rerun()

# Grand total includes discounts (discount_value is negative)
items_total   = sum(it["qty"] * it["unit_price"] for it in st.session_state.fattura_line_items if not it.get("is_discount"))
discount_total = sum(it.get("discount_value", 0.0) for it in st.session_state.fattura_line_items if it.get("is_discount"))
grand_total   = items_total + discount_total
st.markdown(f"### 💰 Total: {currency} {fmt_price(grand_total)}")
if discount_total < 0:
    st.caption(f"Subtotal: {currency} {fmt_price(items_total)} | Discount: {currency} {fmt_price(discount_total)}")

# ── 7. STATUS ──
st.subheader("7. Invoice Status")

# Check if any discount/deduction lines are present
has_discount = any(it.get("is_discount") for it in st.session_state.fattura_line_items)

# Three statuses: not_sent (default), sent, Fattura di anticipo (blue)
# If a discount line is present, Fattura di anticipo is blocked
status_options = {
    "not_sent": "⬜ Not Sent",
    "sent":     "✅ Sent",
    "Fattura di anticipo": "🔵 Fattura di anticipo",
}

if has_discount:
    # Only allow not_sent and sent — this invoice has a deduction so it cannot be an advance invoice
    available_statuses = ["not_sent", "sent"]
    st.caption("ℹ️ *Fattura di anticipo* status is not available when a deduction line is present.")
else:
    available_statuses = list(status_options.keys())

status_choice = st.radio(
    "Status",
    available_statuses,
    format_func=lambda x: status_options[x],
    horizontal=True,
    key="fattura_status"
)

# ── 8. DOCUMENT NAME ──
st.subheader("8. Document Name")
default_name = f"fattura {invoice_number.replace('/','-')} {company}"
doc_name = st.text_input("File name (without .docx)", value=default_name)

# ── GENERATE ──
st.divider()
if st.button("📥 Generate Fattura", type="primary", use_container_width=True):
    if not company:
        st.warning("Please enter a company name.")
        st.stop()
    if not any(it["description"].strip() for it in st.session_state.fattura_line_items):
        st.warning("Please add at least one line item.")
        st.stop()

    # Validate discount lines — must have a linked anticipo selected
    for i, it in enumerate(st.session_state.fattura_line_items):
        if it.get("is_discount") and not it.get("linked_anticipo"):
            st.warning(f"⚠️ Deduction line #{i+1}: please select the advance payment invoice this deduction refers to before generating.")
            st.stop()

    zip_city = f"{zip_code} {city}".strip()
    if region: zip_city += f", {region}"

    try:
        template_path = os.path.join(os.path.dirname(__file__), "fattura_template.docx")
        doc = Document(template_path)
    except Exception as e:
        st.error(f"❌ Template not found: {e}"); st.stop()

    # ── Header paragraphs ──
    header_replacements = {
        "[COMPANY NAME]": company.upper(),
        "[Address]":      address,
        "[Zip] [City], [Region]": zip_city,
        "[Country]":      country,
    }
    for para in doc.paragraphs:
        replace_in_paragraph(para, header_replacements)

    for para in doc.paragraphs:
        full = "".join(r.text for r in para.runs)
        if full.strip() in ["","Messrs."]:
            pass
        else:
            for run in para.runs:
                if run.text.strip():
                    run.bold = False; run.font.name="Verdana"; run.font.size=Pt(10)

    for para in doc.paragraphs:
        if "To the attn. of" in para.text or "All'attenzione" in para.text:
            if include_attn and (salutation or full_name):
                attn_text = f"To the attn. of {salutation} {full_name}".strip().replace("  "," ")
                replace_in_paragraph(para, {"To the attn. of [Sal.] [Full Name]": attn_text})
                for run in para.runs:
                    run.bold=False; run.font.name="Verdana"; run.font.size=Pt(10)
            else:
                delete_para(para)
            break

    # ── Table 0: Invoice details ──
    t0 = doc.tables[0]
    table0_replacements = {
        "[NNN/YY]":                 invoice_number,
        "[DD/MM/YYYY]":             formatted_date,
        "[Partita Iva/VAT Number]": vat_number or "—",
        "[Delivery terms]":         delivery_terms,
    }
    # When status is Fattura di anticipo, relabel the invoice header in the document
    if status_choice == "Fattura di anticipo":
        table0_replacements["INVOICE No.:"]  = "INVOICE FOR ADVANCE PAYMENT No.:"
        table0_replacements["FATTURA N.:"]   = "FATTURA N. (anticipo):"
        table0_replacements["INVOICE NO.:"]  = "INVOICE FOR ADVANCE PAYMENT No.:"
    for row in t0.rows:
        for cell in row.cells:
            replace_in_table_cell(cell, table0_replacements)
            for para in cell.paragraphs:
                full = "".join(r.text for r in para.runs)
                is_invoice = invoice_number in full
                for run in para.runs:
                    run.bold = is_invoice
                    run.font.name="Verdana"; run.font.size=Pt(10)

    # ── Table 1: Payment, bank, delivery ──
    t1 = doc.tables[1]
    set_cell_text(t1.rows[0].cells[0], f"PAYMENT TERMS:\n{payment}",
                  bold=False, font_name="Verdana", font_size=10)
    del_city_region = f"{del_zip} {del_city}".strip()
    if del_region: del_city_region += f", {del_region}"
    del_lines = ["DELIVERY PLACE OF THE GOODS:"]
    if del_company: del_lines.append(del_company)
    if del_address: del_lines.append(del_address)
    if del_city_region: del_lines.append(del_city_region)
    if del_country: del_lines.append(del_country)
    set_cell_text(t1.rows[2].cells[0], "\n".join(del_lines),
                  bold=False, font_name="Verdana", font_size=10)

    # ── Table 2: Products ──
    t2 = doc.tables[2]
    MAX_ROWS = 15
    valid_items = [it for it in st.session_state.fattura_line_items
                   if it["description"].strip()]

    for row_idx in range(1, MAX_ROWS+1):
        row   = t2.rows[row_idx]
        cells = row.cells

        if row_idx-1 < len(valid_items):
            item = valid_items[row_idx-1]
            is_disc = item.get("is_discount", False)

            if is_disc:
                disc_val   = item.get("discount_value", 0.0)
                qty_str    = ""
                price_str  = fmt_price(disc_val)
                set_cell_text(cells[0], "", bold=False)
                desc_cell = cells[1]
                for para in desc_cell.paragraphs:
                    for run in para.runs: run.text=""
                fp = desc_cell.paragraphs[0]
                r_en = fp.add_run(item["description"])
                r_en.bold=False; r_en.font.name="Verdana"; r_en.font.size=Pt(10)
                det_en = item.get("details","").strip()
                if det_en:
                    np2 = copy.deepcopy(fp._p)
                    desc_cell._tc.append(np2)
                    det_para = desc_cell.paragraphs[-1]
                    for run in det_para.runs: run.text=""
                    r_det = det_para.add_run(det_en)
                    r_det.bold=False; r_det.font.name="Verdana"; r_det.font.size=Pt(10)
                set_cell_text(cells[2], "", bold=False)
                set_cell_text(cells[3], "", bold=False)
                set_cell_text(cells[4], currency, bold=False)
                set_cell_text(cells[5], price_str, bold=False)
            else:
                line_total = item["qty"] * item["unit_price"]
                qty_str    = fmt_qty(item["qty"])
                price_str  = fmt_price(item["unit_price"])
                total_str  = fmt_price(line_total)
                set_cell_text(cells[0], qty_str, bold=False)
                desc_cell = cells[1]
                for para in desc_cell.paragraphs:
                    for run in para.runs: run.text=""
                first_para = desc_cell.paragraphs[0]
                r_en = first_para.add_run(item["description"])
                r_en.bold=False; r_en.font.name="Verdana"; r_en.font.size=Pt(10)
                details = item.get("details","").strip()
                if details:
                    new_p2 = copy.deepcopy(first_para._p)
                    desc_cell._tc.append(new_p2)
                    det_para = desc_cell.paragraphs[-1]
                    for run in det_para.runs: run.text=""
                    r_det = det_para.add_run(details)
                    r_det.bold=False; r_det.font.name="Verdana"; r_det.font.size=Pt(10)
                set_cell_text(cells[2], currency, bold=False)
                set_cell_text(cells[3], price_str, bold=False)
                set_cell_text(cells[4], currency, bold=False)
                set_cell_text(cells[5], total_str, bold=False)
        else:
            for cell in cells: set_cell_text(cell,"")
            trPr = row._tr.find(qn('w:trPr'))
            if trPr is None:
                trPr = OxmlElement('w:trPr'); row._tr.insert(0, trPr)
            for old in trPr.findall(qn('w:trHeight')): trPr.remove(old)
            trH = OxmlElement('w:trHeight')
            trH.set(qn('w:val'),'1'); trH.set(qn('w:hRule'),'exact')
            trPr.append(trH)

    # ── Total row ──
    total_row = t2.rows[16]
    tcells    = total_row.cells
    total_label = f"TOTAL AMOUNT \u2013 {delivery_terms} \u2013"
    if vat_exemption:
        total_label += f"\n\n{vat_exemption}"
    set_cell_text(tcells[1], total_label, bold=True)
    set_cell_text(tcells[4], currency, bold=True)
    set_cell_text(tcells[5], fmt_price(grand_total), bold=True)

    # ── HS Code row ──
    hs_row = t2.rows[17]
    set_cell_text(hs_row.cells[1], f"HS code: {hs_code}", bold=False)

    buffer = io.BytesIO()
    doc.save(buffer); buffer.seek(0)

    # Save with chosen status
    fattura_id = save_fattura(
        invoice_number, company, grand_total, currency,
        address, zip_code, city, region, country,
        date_of_reference=selected_date.isoformat(),
        note=note.strip() if note else None,
        payment_terms=payment,
    )

    # Override status if not default
    if fattura_id and status_choice != "not_sent":
        requests.patch(
            f"{SUPABASE_URL}/rest/v1/fatture",
            headers={**HEADERS, "Prefer": "return=minimal"},
            params={"id": f"eq.{fattura_id}"},
            json={"status": status_choice},
        )

    for it in st.session_state.fattura_line_items:
        it["currency"] = currency
    save_fattura_items(fattura_id, st.session_state.fattura_line_items)

    st.success(f"✅ Fattura {invoice_number} ready! Total: {currency} {fmt_price(grand_total)}")
    st.download_button(
        label="📄 Download Word Document", data=buffer,
        file_name=f"{doc_name}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )

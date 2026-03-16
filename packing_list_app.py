import os
import copy
import streamlit as st
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from datetime import date
import io
import requests

st.set_page_config(page_title="Packing List Generator", layout="wide")

def fmt_weight(n):
    """Italian weight format: 1.000,– or 1,25 (2 decimal places)"""
    try:
        f = float(n)
    except (TypeError, ValueError):
        return ""
    cents = round((f % 1) * 100)
    int_str = f"{int(f):,}".replace(",", ".")
    return f"{int_str},–" if cents == 0 else f"{int_str},{cents:02d}"

# ─────────────────────────────────────────────
# PASSWORD GATE
# ─────────────────────────────────────────────
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if not st.session_state.authenticated:
    st.title("🔒 Packing List Generator")
    pwd = st.text_input("Enter passcode to continue:", type="password")
    if st.button("Login"):
        if pwd == "RAINYEAR":
            st.session_state.authenticated = True
            st.rerun()
        else:
            st.error("❌ Wrong passcode.")
    st.stop()

# ─────────────────────────────────────────────
# SUPABASE
# ─────────────────────────────────────────────
SUPABASE_URL = "https://lztrggttkgvgjouofibd.supabase.co"
SUPABASE_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6Imx6dHJnZ3R0a2d2Z2pvdW9maWJkIiwicm9sZSI6ImFub24iLCJpYXQiOjE3NzMyNDAwNzEsImV4cCI6MjA4ODgxNjA3MX0.tbHCQtGW21C2fXCEu2FGwlsXn4kGUWOGoOqjuYyiC7A"
HEADERS = {
    "apikey": SUPABASE_KEY,
    "Authorization": f"Bearer {SUPABASE_KEY}",
    "Content-Type": "application/json"
}

def load_products():
    response = requests.get(
        f"{SUPABASE_URL}/rest/v1/products",
        headers=HEADERS,
        params={"select": "id,description,description_eng,net_weight_kg,dimensions,category",
                "order": "category.asc,created_at.asc"}
    )
    try:
        data = response.json()
        if isinstance(data, list):
            return data
    except:
        pass
    return []

def load_customers():
    response = requests.get(
        f"{SUPABASE_URL}/rest/v1/customers",
        headers=HEADERS,
        params={"select": "id,company_name,contact_name,salutation,address,city,zip,region,country",
                "order": "company_name.asc"}
    )
    try:
        data = response.json()
        if isinstance(data, list):
            return data
    except:
        pass
    return []

def load_fatture():
    response = requests.get(
        f"{SUPABASE_URL}/rest/v1/fatture",
        headers=HEADERS,
        params={"select": "id,invoice_number,client_company,created_at,address,zip,city,region,country",
                "order": "created_at.desc"}
    )
    try:
        data = response.json()
        if isinstance(data, list):
            return data
    except:
        pass
    return []

def load_pl_numbers():
    year_2digit = date.today().strftime('%y')
    r = requests.get(
        f"{SUPABASE_URL}/rest/v1/packing_lists",
        headers=HEADERS,
        params={"select": "pl_number"}
    )
    try:
        d = r.json()
        if isinstance(d, list):
            this_year = [x["pl_number"] for x in d if str(x.get("pl_number", "")).endswith(f"/{year_2digit}")]
            return len(this_year) + 1
    except:
        pass
    return 1

def save_pl_record(pl_number, client_company):
    requests.post(
        f"{SUPABASE_URL}/rest/v1/packing_lists",
        headers={**HEADERS, "Prefer": "return=minimal"},
        json={"pl_number": pl_number, "client_company": client_company}
    )

# ─────────────────────────────────────────────
# DOCX HELPERS  (identical to fattura app)
# ─────────────────────────────────────────────
def set_cell_text(cell, text, bold=False, italic=False, font_name="Verdana", font_size=10):
    tc = cell._tc
    paras = tc.findall(qn('w:p'))
    for extra_p in paras[1:]:
        tc.remove(extra_p)
    first_p = cell.paragraphs[0]
    for run in first_p.runs:
        run.text = ""
        rPr = run._r.find(qn('w:rPr'))
        if rPr is not None:
            run._r.remove(rPr)
    lines = text.split("\n")
    run = first_p.add_run(lines[0])
    run.bold = bold
    run.italic = italic
    run.font.name = font_name
    run.font.size = Pt(font_size)
    for line in lines[1:]:
        br = OxmlElement("w:br")
        run._r.addnext(br)
        run2 = first_p.add_run(line)
        run2.bold = bold
        run2.italic = italic
        run2.font.name = font_name
        run2.font.size = Pt(font_size)
        run = run2

def replace_in_paragraph(para, replacements):
    full_text = "".join(run.text for run in para.runs)
    changed = False
    for key, val in replacements.items():
        if key in full_text:
            full_text = full_text.replace(key, val)
            changed = True
    if changed and para.runs:
        para.runs[0].text = full_text
        for run in para.runs[1:]:
            run.text = ""

def set_para_bold(para, bold):
    for run in para.runs:
        if run.text.strip():
            run.bold = bold

def delete_para(para):
    p = para._p
    p.getparent().remove(p)

# ─────────────────────────────────────────────
# LOAD DATA
# ─────────────────────────────────────────────
if "products_db" not in st.session_state:
    st.session_state.products_db = load_products()
if "customers_db" not in st.session_state:
    st.session_state.customers_db = load_customers()
if "fatture_db" not in st.session_state:
    st.session_state.fatture_db = load_fatture()

PRODUCTS = st.session_state.products_db
CATEGORIES = []
seen_cats = []
for p in PRODUCTS:
    cat = p.get("category") or "Other"
    if cat not in seen_cats:
        seen_cats.append(cat)
        CATEGORIES.append(cat)

PRODUCT_NAMES = ["— custom —"]
PRODUCT_MAP   = {}
for cat in CATEGORIES:
    cat_products = [p for p in PRODUCTS if (p.get("category") or "Other") == cat]
    for p in cat_products:
        eng = p.get("description_eng") or p.get("description", "")
        label = eng[:60] + ("…" if len(eng) > 60 else "")
        PRODUCT_MAP[len(PRODUCT_NAMES)] = p
        PRODUCT_NAMES.append(label)

# ─────────────────────────────────────────────
# SESSION STATE
# ─────────────────────────────────────────────
if "pl_line_items" not in st.session_state:
    st.session_state.pl_line_items = [
        {"product_idx": 0, "description": "", "description_it": "",
         "qty": 1.0, "net_weight": 0.0, "gross_weight": 0.0, "dimensions": ""}
    ]

def add_line():
    st.session_state.pl_line_items.append(
        {"product_idx": 0, "description": "", "description_it": "",
         "qty": 1.0, "net_weight": 0.0, "gross_weight": 0.0, "dimensions": ""}
    )

# ─────────────────────────────────────────────
# UI
# ─────────────────────────────────────────────
st.title("📦 Packing List Generator")

# ── 1. PACKING LIST NUMBER ────────────────────
st.subheader("1. Packing List Number")
year_2digit = date.today().strftime('%y')
next_pl_num = load_pl_numbers()
suggested_pl = f"{next_pl_num:03d}/{year_2digit}"
pl_number = st.text_input("Packing List Number (used in filename only)", value=suggested_pl)

# ── 2. LINK TO FATTURA ────────────────────────
st.subheader("2. Link to Fattura")

fatture = st.session_state.fatture_db
col_fat, col_fat_refresh = st.columns([5, 1])
with col_fat:
    if not fatture:
        st.warning("No fatture found in Supabase.")
        sel_fattura_idx = 0
        fattura_labels  = ["— none —"]
    else:
        fattura_labels = [
            f"{f['invoice_number']} — {f['client_company']} ({f['created_at'][:10]})"
            for f in fatture
        ]
        sel_fattura_idx = st.selectbox(
            "Select Fattura",
            range(len(fattura_labels)),
            format_func=lambda i: fattura_labels[i],
            key="fattura_picker"
        )
with col_fat_refresh:
    st.write("")
    if st.button("🔄", help="Reload fatture"):
        st.session_state.fatture_db = load_fatture()
        st.rerun()

if fatture:
    sel_fattura    = fatture[sel_fattura_idx]
    invoice_number = sel_fattura.get("invoice_number", "")
    client_company = sel_fattura.get("client_company", "")
    fat_date_raw   = sel_fattura.get("created_at", "")
    fat_address    = sel_fattura.get("address", "") or ""
    fat_zip        = sel_fattura.get("zip", "") or ""
    fat_city       = sel_fattura.get("city", "") or ""
    fat_region     = sel_fattura.get("region", "") or ""
    fat_country    = sel_fattura.get("country", "") or ""
    try:
        fattura_date = date.fromisoformat(fat_date_raw[:10]).strftime("%d/%m/%Y")
    except:
        fattura_date = fat_date_raw[:10]
    st.caption(f"📄 Invoice: **{invoice_number}** | Client: **{client_company}** | Date: **{fattura_date}**")
else:
    invoice_number = ""
    client_company = ""
    fattura_date   = date.today().strftime("%d/%m/%Y")
    fat_address = fat_zip = fat_city = fat_region = fat_country = ""

# ── 3. CLIENT ─────────────────────────────────
st.subheader("3. Client")

customers      = st.session_state.customers_db
customer_names = ["— new customer —"] + [
    f"{c.get('company_name', '')} ({c.get('contact_name', '')})" for c in customers
]

# Auto-match customer from fattura's client_company
default_cust_idx = 0
for i, c in enumerate(customers):
    if c.get("company_name", "").strip().lower() == client_company.strip().lower():
        default_cust_idx = i + 1
        break

col_cust, col_refresh = st.columns([5, 1])
with col_cust:
    selected_customer_idx = st.selectbox(
        "Pick existing customer or fill in manually",
        range(len(customer_names)),
        format_func=lambda x: customer_names[x],
        key="cust_picker",
        index=default_cust_idx
    )
with col_refresh:
    st.write("")
    if st.button("🔄", help="Reload customers", key="reload_cust"):
        st.session_state.customers_db = load_customers()
        st.rerun()

if selected_customer_idx > 0:
    cust = customers[selected_customer_idx - 1]
    default_company    = cust.get("company_name", "")
    default_address    = cust.get("address", "") or fat_address
    default_zip        = cust.get("zip", "") or fat_zip
    default_city       = cust.get("city", "") or fat_city
    default_region     = cust.get("region", "") or fat_region or ""
    default_country    = cust.get("country", "") or fat_country
    default_salutation = cust.get("salutation", "Mr.") or "Mr."
    default_full_name  = cust.get("contact_name", "") or ""
else:
    default_company    = client_company
    default_address    = fat_address
    default_zip        = fat_zip
    default_city       = fat_city
    default_region     = fat_region
    default_country    = fat_country
    default_salutation = "Mr."
    default_full_name  = ""

company = st.text_input("Company Name *", value=default_company)
address = st.text_input("Address", value=default_address)
col3, col4, col5 = st.columns(3)
with col3:
    zip_code = st.text_input("Zip", value=default_zip)
with col4:
    city = st.text_input("City", value=default_city)
with col5:
    region = st.text_input("Region", value=default_region, placeholder="(optional)")
country = st.text_input("Country", value=default_country)

include_attn = st.checkbox("Include 'To the attn. of' line?", value=False)
salutation = ""
full_name  = ""
if include_attn:
    col_s, col_n = st.columns([1, 3])
    with col_s:
        sal_opts = ["Mr.", "Ms.", "Dr.", "Messrs."]
        sal_idx  = sal_opts.index(default_salutation) if default_salutation in sal_opts else 0
        salutation = st.selectbox("Salutation", sal_opts, index=sal_idx)
    with col_n:
        full_name = st.text_input("Full Name (optional)", value=default_full_name)

# ── 4. CRATE DIMENSIONS ───────────────────────
st.subheader("4. Crate Dimensions")
crate_dimensions = st.text_input(
    "Crate dimensions (cm)", value="",
    placeholder="e.g. 120 x 80 x 90"
)

# ── 5. LINE ITEMS ─────────────────────────────
st.subheader("5. Line Items")
st.caption("Select from catalogue. Net weight auto-fills from database.")

items_to_remove = []
needs_rerun = False

for i, item in enumerate(st.session_state.pl_line_items):
    with st.container():
        c1, c2, c3, c4 = st.columns([3, 1.5, 1.5, 0.4])
        with c1:
            prod_idx = st.selectbox(
                f"Product #{i+1}",
                range(len(PRODUCT_NAMES)),
                format_func=lambda x: PRODUCT_NAMES[x],
                key=f"pl_prod_{i}",
                index=item["product_idx"]
            )
            if prod_idx != item["product_idx"]:
                item["product_idx"] = prod_idx
                if prod_idx > 0 and prod_idx in PRODUCT_MAP:
                    p = PRODUCT_MAP[prod_idx]
                    item["description"]    = p.get("description_eng") or p.get("description", "")
                    item["description_it"] = p.get("description", "")
                    nw = p.get("net_weight_kg")
                    item["net_weight"]     = float(nw) if nw is not None else 0.0
                    item["gross_weight"]   = item["net_weight"]
                    item["dimensions"]     = p.get("dimensions") or ""
                else:
                    item["description"]    = ""
                    item["description_it"] = ""
                    item["net_weight"]     = 0.0
                    item["gross_weight"]   = 0.0
                    item["dimensions"]     = ""
                needs_rerun = True

            # Show Italian name + dimensions as captions
            if prod_idx > 0 and prod_idx in PRODUCT_MAP:
                p_sel = PRODUCT_MAP[prod_idx]
                it_name = p_sel.get("description", "")
                if it_name:
                    st.caption(f"🇮🇹 {it_name}")
                if p_sel.get("dimensions"):
                    st.caption(f"📐 {p_sel['dimensions']}")

            if prod_idx == 0:
                item["description"] = st.text_input(
                    "Custom Product Name (EN)", value=item.get("description", ""),
                    key=f"pl_desc_{i}")
                item["description_it"] = st.text_input(
                    "Custom Product Name (IT)", value=item.get("description_it", ""),
                    key=f"pl_desc_it_{i}")

        with c2:
            item["qty"] = st.number_input(
                "Qty", min_value=0.0, value=float(item["qty"]),
                step=1.0, format="%.1f", key=f"pl_qty_{i}")

        with c3:
            st.write("**Net Weight (kg)**")
            st.write(fmt_weight(item["net_weight"]) + " kg" if item["net_weight"] else "—")
            item["gross_weight"] = st.number_input(
                "Gross Weight (kg)", min_value=0.0, value=float(item["gross_weight"]),
                step=0.01, format="%.2f", key=f"pl_gw_{i}")

        with c4:
            st.write("")
            st.write("")
            if st.button("🗑", key=f"pl_del_{i}"):
                items_to_remove.append(i)

        line_net   = item["qty"] * item["net_weight"]
        line_gross = item["qty"] * item["gross_weight"]
        st.caption(f"Line net: {fmt_weight(line_net)} kg  |  Line gross: {fmt_weight(line_gross)} kg")
        st.divider()

for i in sorted(items_to_remove, reverse=True):
    st.session_state.pl_line_items.pop(i)
if items_to_remove or needs_rerun:
    st.rerun()

st.button("➕ Add Line Item", on_click=add_line)

valid_items = [it for it in st.session_state.pl_line_items
               if it.get("description", "").strip() and it["qty"] > 0]
total_net   = sum(it["qty"] * it["net_weight"]   for it in valid_items)
total_gross = sum(it["qty"] * it["gross_weight"] for it in valid_items)

col_nw, col_gw = st.columns(2)
with col_nw:
    st.markdown(f"### ⚖️ Total Net: {fmt_weight(total_net)} kg")
with col_gw:
    st.markdown(f"### ⚖️ Total Gross: {fmt_weight(total_gross)} kg")

# ── PREVIEW TABLE ─────────────────────────────
if valid_items:
    st.subheader("Preview")
    preview_data = []
    for it in valid_items:
        preview_data.append({
            "Product": it["description"],
            "Qty": f"{it['qty']:.1f}",
            "Net Wt (kg)": fmt_weight(it["net_weight"]),
            "Gross Wt (kg)": fmt_weight(it["gross_weight"]),
            "Line Net": fmt_weight(it["qty"] * it["net_weight"]),
            "Line Gross": fmt_weight(it["qty"] * it["gross_weight"]),
        })
    st.table(preview_data)

# ── 6. DOCUMENT NAME ──────────────────────────
st.subheader("6. Document Name")
default_name = f"PackingList {pl_number.replace('/', '-')} {company}"
doc_name = st.text_input("File name (without .docx)", value=default_name)

# ── GENERATE ──────────────────────────────────
st.divider()
if st.button("📥 Generate Packing List", type="primary", use_container_width=True):
    if not company:
        st.warning("Please enter a company name.")
    elif not valid_items:
        st.warning("Please add at least one line item.")
    else:
        zip_city = f"{zip_code} {city}".strip()
        if region:
            zip_city += f", {region}"

        try:
            template_path = os.path.join(os.path.dirname(__file__), "packing_list_template.docx")
            doc = Document(template_path)
        except Exception as e:
            st.error(f"❌ Template not found: {e}")
            st.stop()

        # ── Header paragraphs ──
        header_replacements = {
            "[COMPANY NAME]": company.upper(),
            "[Address]":      address,
            "[Zip] [City], [Region]": zip_city,
            "[Country]":      country,
        }
        for para in doc.paragraphs:
            replace_in_paragraph(para, header_replacements)

        # Bold only company, not the rest
        for para in doc.paragraphs:
            full = "".join(r.text for r in para.runs)
            if company.upper() in full:
                set_para_bold(para, True)
            elif full.strip() and full.strip() not in [
                "Messrs.", "PACKING LIST", "Covering the shipment of:",
                "GOODS OF ITALIAN ORIGIN", "All contained in:"
            ]:
                set_para_bold(para, False)

        # Attn line — delete if not needed
        for para in doc.paragraphs:
            if "To the attn. of" in para.text:
                if include_attn and (salutation or full_name):
                    attn_text = f"To the attn. of {salutation} {full_name}".strip().replace("  ", " ")
                    replace_in_paragraph(para, {"To the attn. of [Sal.] [Full Name]": attn_text})
                    set_para_bold(para, False)
                else:
                    delete_para(para)
                break

        # Invoice ref, dimensions, weights — re-fetch paras after possible deletion
        other_replacements = {
            "[NNN/YY]":            invoice_number,
            "[DD/MM/YYYY]":        fattura_date,
            "[dimensions]":        crate_dimensions.strip() if crate_dimensions.strip() else "[dimensions]",
            "[sum of Net Weight]": fmt_weight(total_net),
        }
        for para in doc.paragraphs:
            replace_in_paragraph(para, other_replacements)

        # ── Product table ──
        table    = doc.tables[0]
        MAX_ROWS = 15

        for row_idx in range(1, MAX_ROWS + 1):
            row   = table.rows[row_idx]
            cells = row.cells

            item_idx = row_idx - 1
            if item_idx < len(valid_items):
                item = valid_items[item_idx]

                # Description cell: product name bold, dimensions below if available
                desc_cell  = cells[1]
                for para in desc_cell.paragraphs:
                    for run in para.runs:
                        run.text = ""
                first_para = desc_cell.paragraphs[0]
                r_name = first_para.add_run(item["description"])
                r_name.bold      = True
                r_name.font.name = "Verdana"
                r_name.font.size = Pt(10)
                dims = item.get("dimensions", "")
                if dims:
                    new_p = copy.deepcopy(first_para._p)
                    desc_cell._tc.append(new_p)
                    dim_para = desc_cell.paragraphs[-1]
                    for run in dim_para.runs:
                        run.text = ""
                    r_dim = dim_para.add_run(dims)
                    r_dim.bold      = False
                    r_dim.font.name = "Verdana"
                    r_dim.font.size = Pt(10)

                # Qty: Italian format e.g. "1,0"
                qty_val = item["qty"]
                qty_str = f"{int(qty_val)},0" if qty_val == int(qty_val) else f"{qty_val:.1f}".replace(".", ",")

                set_cell_text(cells[0], qty_str)
                set_cell_text(cells[2], "Kg")
                set_cell_text(cells[3], fmt_weight(item["net_weight"]))
                set_cell_text(cells[4], "Kg")
                set_cell_text(cells[5], fmt_weight(item["gross_weight"]))
            else:
                for cell in cells:
                    set_cell_text(cell, "")
                # Collapse empty rows
                trPr = row._tr.find(qn('w:trPr'))
                if trPr is None:
                    trPr = OxmlElement('w:trPr')
                    row._tr.insert(0, trPr)
                existing_h = trPr.find(qn('w:trHeight'))
                if existing_h is not None:
                    trPr.remove(existing_h)
                trH = OxmlElement('w:trHeight')
                trH.set(qn('w:val'), '1')
                trH.set(qn('w:hRule'), 'exact')
                trPr.append(trH)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        save_pl_record(pl_number, company)

        st.success(f"✅ Packing List {pl_number} ready!")
        st.download_button(
            label="📄 Download Word Document",
            data=buffer,
            file_name=f"{doc_name}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

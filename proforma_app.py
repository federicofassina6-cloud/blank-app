import os
import copy
import streamlit as st
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from datetime import date
import io
import requests

st.set_page_config(page_title="Proforma Generator", layout="wide")

# ─── PASSWORD ───────────────────────────────
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False
if not st.session_state.authenticated:
    st.title("🔒 Proforma Generator")
    pwd = st.text_input("Enter passcode to continue:", type="password")
    if st.button("Login"):
        if pwd == "RAINYEAR":
            st.session_state.authenticated = True
            st.rerun()
        else:
            st.error("❌ Wrong passcode.")
    st.stop()

# ─── LANGUAGE ───────────────────────────────
if "language" not in st.session_state:
    st.session_state.language = None
if st.session_state.language is None:
    st.title("📄 Proforma Invoice Generator")
    st.subheader("Select language / Seleziona lingua")
    c1, c2 = st.columns(2)
    with c1:
        if st.button("🇬🇧  English", use_container_width=True):
            st.session_state.language = "en"; st.rerun()
    with c2:
        if st.button("🇮🇹  Italiano", use_container_width=True):
            st.session_state.language = "it"; st.rerun()
    st.stop()

LANG = st.session_state.language

# ─── SUPABASE ───────────────────────────────
SUPABASE_URL = "https://lztrggttkgvgjouofibd.supabase.co"
SUPABASE_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6Imx6dHJnZ3R0a2d2Z2pvdW9maWJkIiwicm9sZSI6ImFub24iLCJpYXQiOjE3NzMyNDAwNzEsImV4cCI6MjA4ODgxNjA3MX0.tbHCQtGW21C2fXCEu2FGwlsXn4kGUWOGoOqjuYyiC7A"
HDR = {"apikey": SUPABASE_KEY, "Authorization": f"Bearer {SUPABASE_KEY}", "Content-Type": "application/json"}

@st.cache_data(ttl=300)
def load_products():
    try:
        d = requests.get(f"{SUPABASE_URL}/rest/v1/products", headers=HDR,
            params={"select":"id,description,description_eng,unit_price_client,unit_price_reseller,category",
                    "order":"category.asc,created_at.asc"}).json()
        return d if isinstance(d, list) else []
    except: return []

@st.cache_data(ttl=300)
def load_customers():
    try:
        d = requests.get(f"{SUPABASE_URL}/rest/v1/customers", headers=HDR,
            params={"select":"id,company_name,contact_name,salutation,email,phone,address,city,zip,country,notes",
                    "order":"company_name.asc"}).json()
        return d if isinstance(d, list) else []
    except: return []

@st.cache_data(ttl=300)
def load_delivery_terms():
    try:
        d = requests.get(f"{SUPABASE_URL}/rest/v1/delivery_terms", headers=HDR,
            params={"select":"term","order":"created_at.asc"}).json()
        return [x["term"] for x in d] if isinstance(d, list) else []
    except: return []

@st.cache_data(ttl=60)
def load_existing_numbers():
    try:
        d = requests.get(f"{SUPABASE_URL}/rest/v1/fatture_proforma", headers=HDR,
            params={"select":"proforma_number"}).json()
        return [x["proforma_number"] for x in d] if isinstance(d, list) else []
    except: return []

def get_next_number():
    yr = date.today().strftime('%y')
    existing = load_existing_numbers()
    this_yr = [n for n in existing if str(n).endswith(f"/{yr}")]
    return f"{len(this_yr)+1:03d}/{yr}"

def save_proforma(num, company, total, currency, date_of_reference=None):
    # Ensure date_of_reference is a plain string in YYYY-MM-DD format
    if date_of_reference is None:
        date_str = None
    elif isinstance(date_of_reference, str):
        date_str = date_of_reference
    elif hasattr(date_of_reference, "strftime"):
        date_str = date_of_reference.strftime("%Y-%m-%d")
    else:
        date_str = str(date_of_reference)

    payload = {
        "proforma_number": num,
        "client_company": company,
        "total_amount": total,
        "currency": currency,
        "status": "not_sent",
        "date_reference": date_str,
        "date_of_reference": date_str,
    }
    r = requests.post(
        f"{SUPABASE_URL}/rest/v1/fatture_proforma",
        headers={**HDR, "Prefer": "return=representation"},
        json=payload,
    )
    if not r.ok:
        st.warning(f"⚠️ Could not save to Supabase: {r.status_code} — {r.text}")
    else:
        saved = r.json()
        if isinstance(saved, list) and saved:
            saved_date = saved[0].get("date_of_reference")
            if not saved_date:
                st.warning("⚠️ Row saved but date_of_reference is still null. Check Supabase column permissions.")
    load_existing_numbers.clear()

def save_delivery_term(term):
    if term in load_delivery_terms(): return
    requests.post(f"{SUPABASE_URL}/rest/v1/delivery_terms", headers=HDR, json={"term":term})
    load_delivery_terms.clear()

def save_customer(company, contact, salutation, address, city, zip_code, country):
    try:
        ex = requests.get(f"{SUPABASE_URL}/rest/v1/customers", headers=HDR,
            params={"company_name":f"eq.{company}","select":"id"}).json()
        if isinstance(ex, list) and ex: return
    except: pass
    requests.post(f"{SUPABASE_URL}/rest/v1/customers", headers=HDR,
        json={"company_name":company,"contact_name":contact,"salutation":salutation,
              "address":address,"city":city,"zip":zip_code,"country":country})
    load_customers.clear()

# ─── PRICE FORMATTER ────────────────────────
def fmt_it(v: float) -> str:
    if v == int(v):
        return f"{int(v):,}".replace(",",".") + ",\u2013"
    int_p, dec_p = f"{v:,.2f}".split(".")
    return f"{int_p.replace(',','.')},{dec_p}"

# ─── LANGUAGE STRINGS ───────────────────────
if LANG == "en":
    TMPL = "proforma_template_eng.docx"
    TITLE = "📄 Proforma Invoice Generator 🇬🇧"
    TOTAL_TPL = "TOTAL PRICE \u2013 {dt} \u2013"
    PAY_OPTS = ["In advance by T/t transfer","100% by T/T transfer at the order",
                "50% advance, 50% before shipment","30 days from invoice date","Letter of credit at sight"]
    TIME_OPTS = ["2 weeks from payment receipt","3 - 5 weeks from payment receipt",
                 "4 - 6 weeks from payment receipt","6 - 8 weeks from payment receipt","To be confirmed"]
    PACK_OPTS = ["Included, for air shipment","Included with fumigated wooden crate, for air shipment",
                 "Included with carton box","Not included"]
    SHIP_OPTS = ["By express courier","By air","By sea","By road","To be arranged by customer"]
    L = dict(date="Date",client="2. Client",pick="Pick existing customer or fill in manually",
             reload="🔄",sal="Salutation",contact="Contact Full Name (optional)",
             company="Company Name *",addr="Address",zip="Zip",city="City",region="Region",country="Country",
             cur_sec="3. Currency & Price Type",cur="Currency (ISO)",ptype="Price type",
             lines="4. Line Items",lines_cap="Select from catalogue.",prod="Product #{i}",
             details="Specs / Description (optional)",qty="Qty",uprice="Unit Price ({c})",
             rm="🗑",addline="➕ Add Line Item",terms="5. Terms & Conditions",hs="HS Code",
             pay="Payment",dt="Delivery Terms",dtime="Delivery Time",pack="Packing",ship="Shipment",
             savedt="💾 Save this term",docname="6. Document Name",fname="File name (without .docx)",
             gen="📥 Generate Proforma Invoice",wcomp="Please enter a company name.",
             witems="Please add at least one line item.",
             ok="✅ Proforma {n} ready! Total: {c} {t}",dl="📄 Download Word Document",
             cust="— custom —",newcust="— new customer —",cli="Cliente",riv="Rivenditore",
             lswitch="🇮🇹 Switch to Italian",attn="Include 'To the attention of' line?",
             nlabel="Proforma Number",nhint="Suggested — you can change it",
             nwarn="⚠️ Number outside normal sequence.",ndup="❌ This number already exists.")
else:
    TMPL = "proforma_template_ita.docx"
    TITLE = "📄 Generatore Fattura Proforma 🇮🇹"
    TOTAL_TPL = "TOTALE \u2013 {dt} \u2013"
    PAY_OPTS = ["Anticipato tramite bonifico","100% bonifico all'ordine",
                "50% anticipo, 50% prima spedizione","30 giorni dalla fattura","Lettera di credito a vista"]
    TIME_OPTS = ["2 settimane dal pagamento","3-5 settimane dal pagamento",
                 "4-6 settimane dal pagamento","6-8 settimane dal pagamento","Da confermare"]
    PACK_OPTS = ["Incluso, spedizione aerea","Incluso in cassa fumigata, spedizione aerea",
                 "Incluso in scatola cartone","Non incluso"]
    SHIP_OPTS = ["Corriere espresso","Via aerea","Via mare","Via strada","A cura del cliente"]
    L = dict(date="Data",client="2. Cliente",pick="Seleziona cliente o compila manualmente",
             reload="🔄",sal="Titolo",contact="Nome contatto (opzionale)",
             company="Ragione sociale *",addr="Indirizzo",zip="CAP",city="Città",region="Provincia",country="Paese",
             cur_sec="3. Valuta e tipo prezzo",cur="Valuta (ISO)",ptype="Tipo prezzo",
             lines="4. Articoli",lines_cap="Seleziona dal catalogo.",prod="Prodotto #{i}",
             details="Descrizione / Specifiche (opzionale)",qty="Q.tà",uprice="Prezzo unitario ({c})",
             rm="🗑",addline="➕ Aggiungi articolo",terms="5. Condizioni generali",hs="Codice HS",
             pay="Pagamento",dt="Resa",dtime="Consegna",pack="Imballo",ship="Spedizione",
             savedt="💾 Salva resa",docname="6. Nome documento",fname="Nome file (senza .docx)",
             gen="📥 Genera Fattura Proforma",wcomp="Inserire la ragione sociale.",
             witems="Aggiungere almeno un articolo.",
             ok="✅ Proforma {n} pronta! Totale: {c} {t}",dl="📄 Scarica documento Word",
             cust="— personalizzato —",newcust="— nuovo cliente —",cli="Cliente",riv="Rivenditore",
             lswitch="🇬🇧 Switch to English",attn="Includere riga 'All'attenzione di'?",
             nlabel="Numero Proforma",nhint="Progressivo suggerito — modificabile",
             nwarn="⚠️ Numero fuori sequenza.",ndup="❌ Numero già esistente.")

HS = ["8453.9000","8453.1000","8466.9195","8464.2019","8451.9000","8451.8030"]
CURS = ["EUR","USD","GBP","CHF","CNY","RUB",L["cust"]]

# ─── LOAD DATA ──────────────────────────────
if "products_db" not in st.session_state:
    st.session_state.products_db = load_products()
if "customers_db" not in st.session_state:
    st.session_state.customers_db = load_customers()
if "dt_db" not in st.session_state:
    st.session_state.dt_db = load_delivery_terms()

PRODS = st.session_state.products_db

# Flat product list — NO category separators
PNAMES = ["— select product —"]
PMAP   = {}
for p in PRODS:
    dk    = "description" if LANG == "it" else "description_eng"
    label = (p.get(dk) or p.get("description") or "")
    label = label[:60] + ("…" if len(label) > 60 else "")
    PMAP[len(PNAMES)] = p
    PNAMES.append(label)

# ─── DOCX HELPERS ───────────────────────────
def replace_para(para, reps):
    for key, val in reps.items():
        full = "".join(r.text for r in para.runs)
        if key not in full: continue
        keeper = next((r for r in para.runs if key in r.text), None)
        if not keeper: keeper = next((r for r in para.runs if r.bold), para.runs[-1] if para.runs else None)
        new_text = full.replace(key, val)
        if para.runs:
            para.runs[0].text = new_text
            if keeper and keeper != para.runs[0]:
                para.runs[0].bold   = keeper.bold
                para.runs[0].italic = keeper.italic
                if keeper.font.name: para.runs[0].font.name = keeper.font.name
                if keeper.font.size: para.runs[0].font.size = keeper.font.size
            for r in para.runs[1:]: r.text = ""

def set_run(para, text, bold=False, fn="Verdana", fs=10):
    para.clear()
    r = para.add_run(text)
    r.bold = bold; r.font.name = fn; r.font.size = Pt(fs)

def set_cell(cell, text, bold=False, fn="Verdana", fs=10):
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
            rp = r._r.find(qn('w:rPr'))
            if rp is not None: r._r.remove(rp)
    r = cell.paragraphs[0].add_run(text)
    r.bold = bold; r.font.name = fn; r.font.size = Pt(fs)

def collapse_para(para):
    para.clear()
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._p.insert(0, pPr)
    for old in pPr.findall(qn('w:spacing')): pPr.remove(old)
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '0'); sp.set(qn('w:after'), '0')
    sp.set(qn('w:line'), '120'); sp.set(qn('w:lineRule'), 'exact')
    pPr.append(sp)

def delete_para(para):
    p = para._p
    p.getparent().remove(p)

# ─── SESSION STATE ───────────────────────────
if "line_items" not in st.session_state:
    st.session_state.line_items = [
        {"product_idx":0,"description":"","details":"","qty":1.0,"unit_price":0.0}]

def add_line():
    st.session_state.line_items.append(
        {"product_idx":0,"description":"","details":"","qty":1.0,"unit_price":0.0})

# ─── UI ─────────────────────────────────────
col_t, col_l = st.columns([5,1])
with col_t: st.title(TITLE)
with col_l:
    st.write("")
    if st.button(L["lswitch"]):
        st.session_state.language = None
        st.session_state.line_items = []
        st.rerun()

# 1. DATE & NUMBER
st.subheader(f"1. {L['date']} & {L['nlabel']}")
cd1, cd2 = st.columns(2)
with cd1:
    sel_date_raw = st.date_input(L["date"], value=date.today(), format="DD/MM/YYYY")
    # Normalise: date_input can return a tuple on some Streamlit versions
    if isinstance(sel_date_raw, (list, tuple)):
        sel_date = sel_date_raw[0] if sel_date_raw else date.today()
    else:
        sel_date = sel_date_raw
    # Store in session state so it survives reruns
    st.session_state["sel_date"] = sel_date

with cd2:
    yr2 = sel_date.strftime('%y')
    suggested = get_next_number()
    existing_nums = load_existing_numbers()
    pnum = st.text_input(L["nlabel"], value=suggested, help=L["nhint"]).strip()
    number_ok = True
    if pnum in existing_nums:
        st.error(L["ndup"])
        number_ok = False
    elif pnum:
        try:
            seq = int(pnum.split("/")[0])
            exp = int(suggested.split("/")[0])
            if seq != exp: st.warning(L["nwarn"])
        except: pass

fmt_date = sel_date.strftime('%d/%m/') + "\u2019" + yr2

# 2. CLIENT
st.subheader(L["client"])
custs = st.session_state.customers_db
cnames = [L["newcust"]] + [f"{c.get('company_name','')} ({c.get('contact_name','')})" for c in custs]
cc1, cc2 = st.columns([5,1])
with cc1:
    cidx = st.selectbox(L["pick"], range(len(cnames)), format_func=lambda x: cnames[x], key="cpick")
with cc2:
    st.write("")
    if st.button(L["reload"]):
        load_customers.clear()
        st.session_state.customers_db = load_customers()
        st.rerun()

if cidx > 0:
    c = custs[cidx-1]
    dsal = c.get("salutation","Mr.") or "Mr."
    dsal = dsal if dsal in ["Mr.","Ms.","Dr.","Messrs."] else "Mr."
    dfn=c.get("contact_name",""); dco=c.get("company_name","")
    dad=c.get("address",""); dzip=c.get("zip",""); dcity=c.get("city",""); dctry=c.get("country","")
else:
    dsal="Mr."; dfn=dco=dad=dzip=dcity=dctry=""

include_attn = st.checkbox(L["attn"], value=True)
if include_attn:
    ca1, ca2 = st.columns([1,3])
    with ca1: sal = st.selectbox(L["sal"],["Mr.","Ms.","Dr.","Messrs."],
                                  index=["Mr.","Ms.","Dr.","Messrs."].index(dsal))
    with ca2: full_name = st.text_input(L["contact"], value=dfn)
else:
    sal=""; full_name=""

company = st.text_input(L["company"], value=dco)
address = st.text_input(L["addr"], value=dad)
cc3,cc4,cc5 = st.columns(3)
with cc3: zip_code = st.text_input(L["zip"], value=dzip)
with cc4: city = st.text_input(L["city"], value=dcity)
with cc5: region = st.text_input(L["region"], value="")
country = st.text_input(L["country"], value=dctry)

# 3. CURRENCY & PRICE TYPE
st.subheader(L["cur_sec"])
cp1, cp2 = st.columns(2)
with cp1:
    cur_choice = st.selectbox(L["cur"], CURS)
    currency = st.text_input("ISO code", placeholder="e.g. AED") if cur_choice == L["cust"] else cur_choice
with cp2:
    gpt = st.radio(L["ptype"], [L["cli"], L["riv"]], horizontal=True, key="gpt")
    if st.session_state.get("_lpt") != gpt:
        st.session_state["_lpt"] = gpt
        for it in st.session_state.line_items:
            if it.get("product_idx",0) > 0 and it.get("product_idx") in PMAP:
                p = PMAP[it["product_idx"]]
                pc = float(p.get("unit_price_client") or 0)
                pr = float(p.get("unit_price_reseller") or 0)
                it["unit_price"] = pc if gpt == L["cli"] else pr
        st.rerun()

# 4. LINE ITEMS
st.subheader(L["lines"])
st.caption(L["lines_cap"])
to_rm = []; needs_rerun = False
for i, item in enumerate(st.session_state.line_items):
    with st.container():
        lc1,lc2,lc3,lc4 = st.columns([3,1.5,1.5,0.4])
        with lc1:
            pidx = st.selectbox(L["prod"].replace("{i}",str(i+1)),
                range(len(PNAMES)), format_func=lambda x: PNAMES[x],
                key=f"p_{i}", index=item["product_idx"])
            if pidx != item["product_idx"]:
                item["product_idx"] = pidx
                if pidx > 0 and pidx in PMAP:
                    p = PMAP[pidx]
                    dk = "description" if LANG=="it" else "description_eng"
                    item["description"] = p.get(dk) or p.get("description") or ""
                    pc = float(p.get("unit_price_client") or 0)
                    pr = float(p.get("unit_price_reseller") or 0)
                    item["unit_price"] = pc if gpt == L["cli"] else pr
                else:
                    item["description"]=""; item["unit_price"]=0.0
                needs_rerun = True
            if pidx > 0 and pidx in PMAP:
                pp = PMAP[pidx]
                if pp.get("description"): st.caption(f"🇮🇹 {pp['description']}")
                if pp.get("description_eng"): st.caption(f"🇬🇧 {pp['description_eng']}")
            item["details"] = st.text_input(L["details"], value=item.get("details",""), key=f"d_{i}")
        with lc2:
            item["qty"] = st.number_input(L["qty"], min_value=0.0, value=float(item["qty"]),
                step=1.0, format="%.1f", key=f"q_{i}")
        with lc3:
            st.write(f"**{L['uprice'].format(c=currency)}**")
            st.write(fmt_it(item["unit_price"]))
        with lc4:
            st.write(""); st.write("")
            if st.button(L["rm"], key=f"r_{i}"): to_rm.append(i)
        lt = item["qty"] * item["unit_price"]
        st.caption(f"Line total: {currency} {fmt_it(lt)}")
        st.divider()

for i in sorted(to_rm, reverse=True): st.session_state.line_items.pop(i)
if to_rm or needs_rerun: st.rerun()

st.button(L["addline"], on_click=add_line)
grand_total = sum(it["qty"]*it["unit_price"] for it in st.session_state.line_items)
st.markdown(f"### 💰 Total: {currency} {fmt_it(grand_total)}")

# 5. TERMS
st.subheader(L["terms"])
DT_OPTS = st.session_state.dt_db
tc1, tc2 = st.columns(2)
with tc1:
    hs = st.selectbox(L["hs"], HS+[L["cust"]])
    if hs == L["cust"]: hs = st.text_input("Custom HS")
    pay = st.selectbox(L["pay"], PAY_OPTS+[L["cust"]])
    if pay == L["cust"]: pay = st.text_input("Custom payment")
    dt = st.selectbox(L["dt"], DT_OPTS+[L["cust"]])
    if dt == L["cust"]:
        dt = st.text_input("Custom delivery term")
        if dt and dt not in DT_OPTS:
            if st.button(L["savedt"]):
                save_delivery_term(dt)
                st.session_state.dt_db = load_delivery_terms()
                st.success(f"✅ '{dt}' saved!"); st.rerun()
    dtime = st.selectbox(L["dtime"], TIME_OPTS+[L["cust"]])
    if dtime == L["cust"]: dtime = st.text_input("Custom time")
with tc2:
    pack = st.selectbox(L["pack"], PACK_OPTS+[L["cust"]])
    if pack == L["cust"]: pack = st.text_input("Custom packing")
    ship = st.selectbox(L["ship"], SHIP_OPTS+[L["cust"]])
    if ship == L["cust"]: ship = st.text_input("Custom shipment")

# 6. DOC NAME
st.subheader(L["docname"])
doc_name = st.text_input(L["fname"], value=f"proforma {pnum.replace('/','.')} {company}")

# GENERATE
st.divider()
if st.button(L["gen"], type="primary", use_container_width=True, disabled=not number_ok):
    if not company:
        st.warning(L["wcomp"])
    elif not any(it["description"].strip() for it in st.session_state.line_items):
        st.warning(L["witems"])
    else:
        # Retrieve the date safely from session state
        final_date = st.session_state.get("sel_date", date.today())
        if isinstance(final_date, (list, tuple)):
            final_date = final_date[0] if final_date else date.today()
        date_str = final_date.strftime("%Y-%m-%d")

        zip_city = f"{zip_code} {city}".strip()
        if region: zip_city += f", {region}"

        reps = {
            f"Schio, [DD/MM/\u2019YY]": f"Schio, {fmt_date}",
            f"[DD/MM/\u2019YY]": fmt_date,
            "[COMPANY NAME]": company,
            "[Address]": address,
            "[Zip] [City], [Region]": zip_city,
            "[Country]": country,
            "Mr./Ms. [Full Name]": f"{sal} {full_name}".strip() if include_attn else "",
            "[Sal.]": sal,
            "[Full Name]": full_name,
            "[NNN/YY]": pnum,
        }

        try:
            tpath = os.path.join(os.path.dirname(__file__), TMPL)
            doc = Document(tpath)
        except Exception as e:
            st.error(f"❌ Template not found: {e}"); st.stop()

        for para in doc.paragraphs:
            replace_para(para, reps)

        for para in doc.paragraphs:
            full = "".join(r.text for r in para.runs)

            if para == doc.paragraphs[0]:
                para.clear()
                r1 = para.add_run("Schio, ")
                r1.bold=False; r1.font.name="Verdana"; r1.font.size=Pt(10)
                r2 = para.add_run(fmt_date)
                r2.bold=False; r2.font.name="Verdana"; r2.font.size=Pt(10)
                continue

            if "PROFORMA INVOICE NO" in full or "FATTURA PROFORMA N" in full or "PROFORMA" in full.upper() and "N" in full.upper() and pnum in full:
                for r in para.runs:
                    r.bold=True; r.font.name="Verdana"; r.font.size=Pt(10)
                continue

            if para.text.strip().startswith("To the attn."):
                if include_attn and (sal or full_name):
                    para.clear()
                    attn_text = f"To the attn. of {sal or ''} {full_name or ''}".strip().replace("  ", " ")
                    run = para.add_run(attn_text)
                    run.bold = False; run.font.name = "Verdana"; run.font.size = Pt(10)
                else:
                    delete_para(para)
                continue

            if company and company in full:
                set_run(para, company, bold=True)
                continue

            for r in para.runs:
                r.bold=False; r.font.name="Verdana"; r.font.size=Pt(10)

        tbl = doc.tables[0]
        MAX = 15
        valid = [it for it in st.session_state.line_items if it["description"].strip()]

        for ri in range(1, MAX+1):
            row = tbl.rows[ri]
            cells = row.cells
            if ri-1 < len(valid):
                it = valid[ri-1]
                lt = it["qty"] * it["unit_price"]
                qty_s = f"{it['qty']:,.1f}".replace(",","X").replace(".",",").replace("X",".")
                set_cell(cells[0], str(ri*10))
                dc = cells[1]
                for p in dc.paragraphs:
                    for r in p.runs:
                        r.text=""
                        rp = r._r.find(qn('w:rPr'))
                        if rp is not None: r._r.remove(rp)
                fp = dc.paragraphs[0]
                nr = fp.add_run(it["description"])
                nr.bold=True; nr.font.name="Verdana"; nr.font.size=Pt(10)
                det = it.get("details","").strip()
                if det:
                    np2 = copy.deepcopy(fp._p)
                    dc._tc.append(np2)
                    sp = dc.paragraphs[-1]
                    for r in sp.runs: r.text=""
                    dr = sp.add_run(det)
                    dr.bold=False; dr.font.name="Verdana"; dr.font.size=Pt(10)
                set_cell(cells[2], qty_s)
                set_cell(cells[3], fmt_it(it["unit_price"]))
                set_cell(cells[4], currency)
                set_cell(cells[5], fmt_it(lt))
            else:
                for c in cells: set_cell(c, "")
                trPr = row._tr.find(qn('w:trPr'))
                if trPr is None:
                    trPr = OxmlElement('w:trPr'); row._tr.insert(0, trPr)
                for old in trPr.findall(qn('w:trHeight')): trPr.remove(old)
                trH = OxmlElement('w:trHeight')
                trH.set(qn('w:val'),'1'); trH.set(qn('w:hRule'),'exact')
                trPr.append(trH)

        tr = tbl.rows[MAX+1]
        tc = tr.cells
        set_cell(tc[0], TOTAL_TPL.format(dt=dt), bold=True)
        set_cell(tc[4], currency, bold=True)
        set_cell(tc[5], fmt_it(grand_total), bold=True)

        tt = doc.tables[1]
        for ri2, val in {0:hs,1:pay,4:dt,5:dtime,6:pack,7:ship}.items():
            if ri2 < len(tt.rows):
                set_cell(tt.rows[ri2].cells[1], val)

        buf = io.BytesIO()
        doc.save(buf); buf.seek(0)

        save_proforma(pnum, company, grand_total, currency, date_of_reference=date_str)

        if company.strip():
            save_customer(company, full_name, sal, address, city, zip_code, country)
            load_customers.clear()
            st.session_state.customers_db = load_customers()

        st.success(L["ok"].format(n=pnum, c=currency, t=fmt_it(grand_total)))
        st.download_button(label=L["dl"], data=buf, file_name=f"{doc_name}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True)

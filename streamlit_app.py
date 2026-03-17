import os
import copy
import streamlit as st
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from datetime import date
import io
import requests

st.set_page_config(page_title="Offerta Generator", layout="wide")

# ─────────────────────────────────────────────
# PASSWORD GATE
# ─────────────────────────────────────────────
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if not st.session_state.authenticated:
    st.title("🔒 Offerta Generator")
    pwd = st.text_input("Enter passcode to continue:", type="password")
    if st.button("Login"):
        if pwd == "RAINYEAR":
            st.session_state.authenticated = True
            st.rerun()
        else:
            st.error("❌ Wrong passcode.")
    st.stop()

# ─────────────────────────────────────────────
# LANGUAGE SELECTION
# ─────────────────────────────────────────────
if "language" not in st.session_state:
    st.session_state.language = None

if st.session_state.language is None:
    st.title("📄 Offerta Generator")
    st.subheader("Select language / Seleziona lingua")
    col_en, col_it = st.columns(2)
    with col_en:
        if st.button("🇬🇧  English", use_container_width=True):
            st.session_state.language = "en"
            st.rerun()
    with col_it:
        if st.button("🇮🇹  Italiano", use_container_width=True):
            st.session_state.language = "it"
            st.rerun()
    st.stop()

LANG = st.session_state.language

# ─────────────────────────────────────────────
# SUPABASE
# ─────────────────────────────────────────────
SUPABASE_URL = "https://lztrggttkgvgjouofibd.supabase.co"
SUPABASE_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6Imx6dHJnZ3R0a2d2Z2pvdW9maWJkIiwicm9sZSI6ImFub24iLCJpYXQiOjE3NzMyNDAwNzEsImV4cCI6MjA4ODgxNjA3MX0.tbHCQtGW21C2fXCEu2FGwlsXn4kGUWOGoOqjuYyiC7A"
HEADERS = {
    "apikey": SUPABASE_KEY,
    "Authorization": f"Bearer {SUPABASE_KEY}",
    "Content-Type": "application/json"
}

@st.cache_data(ttl=300)
def load_products():
    r = requests.get(f"{SUPABASE_URL}/rest/v1/products", headers=HEADERS,
        params={"select": "id,description,description_eng,unit_price_client,unit_price_reseller,category",
                "order": "category.asc,created_at.asc"})
    try:
        d = r.json()
        return d if isinstance(d, list) else []
    except:
        return []

@st.cache_data(ttl=300)
def load_customers():
    r = requests.get(f"{SUPABASE_URL}/rest/v1/customers", headers=HEADERS,
        params={"select": "id,company_name,contact_name,salutation,email,phone,address,city,zip,country,notes",
                "order": "company_name.asc"})
    try:
        d = r.json()
        return d if isinstance(d, list) else []
    except:
        return []

@st.cache_data(ttl=300)
def load_delivery_terms():
    r = requests.get(f"{SUPABASE_URL}/rest/v1/delivery_terms", headers=HEADERS,
        params={"select": "term", "order": "created_at.asc"})
    try:
        d = r.json()
        return [x["term"] for x in d] if isinstance(d, list) else []
    except:
        return []

@st.cache_data(ttl=60)
def load_existing_offerta_numbers():
    r = requests.get(f"{SUPABASE_URL}/rest/v1/offerte", headers=HEADERS,
        params={"select": "offer_number"})
    try:
        d = r.json()
        return [x["offer_number"] for x in d if x.get("offer_number")] if isinstance(d, list) else []
    except:
        return []

def get_next_offerta_number():
    year_2digit = date.today().strftime('%y')
    existing = load_existing_offerta_numbers()
    # Only count OF-prefixed numbers for this year
    this_year = [n for n in existing if str(n).startswith("OF") and str(n).endswith(f"/{year_2digit}")]
    return f"OF{len(this_year) + 1:03d}/{year_2digit}"

def save_offerta(offerta_number, client_company, total_amount, currency, date_of_reference=None):
    payload = {
        "offer_number": offerta_number,
        "client_company": client_company,
        "date_of_reference": date_of_reference,
    }
    r = requests.post(
        f"{SUPABASE_URL}/rest/v1/offerte",
        headers={**HEADERS, "Prefer": "return=minimal"},
        json=payload,
    )
    if not r.ok:
        st.warning(f"⚠️ Could not save to Supabase: {r.status_code} {r.text}")
    load_existing_offerta_numbers.clear()

def save_delivery_term(term):
    existing = load_delivery_terms()
    if term in existing:
        return
    requests.post(f"{SUPABASE_URL}/rest/v1/delivery_terms", headers=HEADERS, json={"term": term})
    load_delivery_terms.clear()

def save_customer(company_name, contact_name, salutation, email, phone, address, city, zip_code, country, notes):
    r = requests.get(f"{SUPABASE_URL}/rest/v1/customers", headers=HEADERS,
        params={"company_name": f"eq.{company_name}", "select": "id"})
    try:
        if isinstance(r.json(), list) and len(r.json()) > 0:
            return
    except:
        pass
    requests.post(f"{SUPABASE_URL}/rest/v1/customers", headers=HEADERS,
        json={"company_name": company_name, "contact_name": contact_name,
              "salutation": salutation, "email": email, "phone": phone,
              "address": address, "city": city, "zip": zip_code,
              "country": country, "notes": notes})
    load_customers.clear()

# ─────────────────────────────────────────────
# ITALIAN PRICE FORMATTER
# ─────────────────────────────────────────────
def fmt_price_it(value: float) -> str:
    if value == int(value):
        int_part = f"{int(value):,}".replace(",", ".")
        return f"{int_part},\u2013"
    else:
        formatted = f"{value:,.2f}"
        int_part, dec_part = formatted.split(".")
        int_part = int_part.replace(",", ".")
        return f"{int_part},{dec_part}"

# ─────────────────────────────────────────────
# LANGUAGE STRINGS & OPTIONS
# ─────────────────────────────────────────────
if LANG == "en":
    TEMPLATE_FILE = "offerta_template_eng.docx"
    TITLE         = "📄 Offer Generator 🇬🇧"
    TOTAL_LABEL_TPL = "TOTAL PRICE \u2013 {dt} \u2013"
    PAYMENT_OPTIONS = [
        "In advance by T/t transfer",
        "100% by T/T transfer at the order",
        "50% advance, 50% before shipment",
        "30 days from invoice date",
        "Letter of credit at sight",
    ]
    DELIVERY_TIME_OPTIONS = [
        "2 weeks from payment receipt",
        "3 - 5 weeks from payment receipt",
        "4 - 6 weeks from payment receipt",
        "6 - 8 weeks from payment receipt",
        "To be confirmed",
    ]
    PACKING_OPTIONS = [
        "Included, for air shipment",
        "Included with fumigated wooden crate, for air shipment",
        "Included with carton box",
        "Not included",
    ]
    SHIPMENT_OPTIONS = [
        "By express courier", "By air", "By sea", "By road", "To be arranged by customer",
    ]
    LBL = {
        "date": "Date", "client": "2. Client",
        "pick_cust": "Pick existing customer or fill in manually below",
        "reload": "🔄", "salutation": "Salutation", "contact": "Contact Full Name (optional)",
        "company": "Company Name *", "address": "Address", "zip": "Zip", "city": "City",
        "region": "Region", "country": "Country", "currency": "3. Currency & Price Type",
        "cur_lbl": "Currency (ISO)", "price_type": "Price type (applies to all products)",
        "lines": "4. Line Items", "lines_cap": "Select from catalogue.",
        "prod": "Product #{i} (bold in document)",
        "details": "Description / Specs (optional)", "qty": "Qty",
        "unit_price": "Unit Price ({cur})", "remove": "🗑", "add_line": "➕ Add Line Item",
        "terms": "5. Terms & Conditions", "hs": "HS Code", "payment": "Payment",
        "del_terms": "Delivery Terms", "del_time": "Delivery Time", "packing": "Packing",
        "shipment": "Shipment", "save_dt": "💾 Save this delivery term",
        "doc_name": "6. Document Name", "file_name": "File name (without .docx)",
        "generate": "📥 Generate Offer",
        "warn_company": "Please enter a company name.",
        "warn_items": "Please add at least one line item.",
        "success": "✅ Offerta {num} ready! Total: {cur} {total}",
        "download": "📄 Download Word Document",
        "custom": "— custom —", "new_cust": "— new customer —",
        "cliente": "Cliente", "rivenditore": "Rivenditore",
        "lang_switch": "🇮🇹 Switch to Italian",
        "attn_toggle": "Include 'To the attention of' line?",
        "number_label": "Offer Number",
        "number_hint": "Suggested next offer number — you can change it",
        "number_warn": "⚠️ This number is outside the normal sequence. Continue anyway?",
        "number_dup": "❌ This offer number already exists. Please choose a different one.",
    }
else:
    TEMPLATE_FILE = "offerta_template_ita.docx"
    TITLE         = "📄 Generatore Offerta 🇮🇹"
    TOTAL_LABEL_TPL = "TOTALE \u2013 {dt} \u2013"
    PAYMENT_OPTIONS = [
        "Anticipato tramite bonifico bancario",
        "100% bonifico bancario all'ordine",
        "50% anticipo, 50% prima della spedizione",
        "30 giorni dalla data fattura",
        "Lettera di credito a vista",
    ]
    DELIVERY_TIME_OPTIONS = [
        "2 settimane dal ricevimento pagamento",
        "3 - 5 settimane dal ricevimento pagamento",
        "4 - 6 settimane dal ricevimento pagamento",
        "6 - 8 settimane dal ricevimento pagamento",
        "Da confermare",
    ]
    PACKING_OPTIONS = [
        "Incluso, per spedizione aerea",
        "Incluso in cassa di legno fumigata, per spedizione aerea",
        "Incluso in scatola di cartone",
        "Non incluso",
    ]
    SHIPMENT_OPTIONS = [
        "Corriere espresso", "Via aerea", "Via mare", "Via strada", "A cura del cliente",
    ]
    LBL = {
        "date": "Data", "client": "2. Cliente",
        "pick_cust": "Seleziona cliente o compila manualmente",
        "reload": "🔄", "salutation": "Titolo", "contact": "Nome completo contatto (opzionale)",
        "company": "Ragione sociale *", "address": "Indirizzo", "zip": "CAP", "city": "Città",
        "region": "Provincia", "country": "Paese", "currency": "3. Valuta e tipo prezzo",
        "cur_lbl": "Valuta (ISO)", "price_type": "Tipo prezzo (valido per tutti i prodotti)",
        "lines": "4. Articoli", "lines_cap": "Seleziona dal catalogo.",
        "prod": "Prodotto #{i} (grassetto nel documento)",
        "details": "Descrizione / Specifiche (opzionale)", "qty": "Q.tà",
        "unit_price": "Prezzo unitario ({cur})", "remove": "🗑", "add_line": "➕ Aggiungi articolo",
        "terms": "5. Condizioni generali", "hs": "Codice HS", "payment": "Pagamento",
        "del_terms": "Resa", "del_time": "Consegna", "packing": "Imballo",
        "shipment": "Spedizione", "save_dt": "💾 Salva questa resa",
        "doc_name": "6. Nome documento", "file_name": "Nome file (senza .docx)",
        "generate": "📥 Genera Offerta",
        "warn_company": "Inserire la ragione sociale.",
        "warn_items": "Aggiungere almeno un articolo.",
        "success": "✅ Offerta {num} pronta! Totale: {cur} {total}",
        "download": "📄 Scarica documento Word",
        "custom": "— personalizzato —", "new_cust": "— nuovo cliente —",
        "cliente": "Cliente", "rivenditore": "Rivenditore",
        "lang_switch": "🇬🇧 Switch to English",
        "attn_toggle": "Includere riga 'All'attenzione di'?",
        "number_label": "Numero Offerta",
        "number_hint": "Numero offerta progressivo suggerito — puoi modificarlo",
        "number_warn": "⚠️ Questo numero è fuori dalla sequenza normale. Continuare?",
        "number_dup": "❌ Questo numero offerta esiste già. Sceglierne uno diverso.",
    }

HS_CODES = ["8453.9000","8453.1000","8466.9195","8464.2019","8451.9000","8451.8030"]
CURRENCIES = ["EUR", "USD", "GBP", "CHF", "CNY", "RUB", LBL["custom"]]

# ─────────────────────────────────────────────
# LOAD DATA
# ─────────────────────────────────────────────
if "products_db" not in st.session_state:
    st.session_state.products_db = load_products()
if "customers_db" not in st.session_state:
    st.session_state.customers_db = load_customers()
if "delivery_terms_db" not in st.session_state:
    st.session_state.delivery_terms_db = load_delivery_terms()

PRODUCTS = st.session_state.products_db
CATEGORIES = []
seen_cats = []
for p in PRODUCTS:
    cat = p.get("category") or "Other"
    if cat not in seen_cats:
        seen_cats.append(cat)
        CATEGORIES.append(cat)

PRODUCT_NAMES = ["— select product —"]
PRODUCT_MAP   = {}
for cat in CATEGORIES:
    cat_products = [p for p in PRODUCTS if (p.get("category") or "Other") == cat]
    for p in cat_products:
        desc_key = "description" if LANG == "it" else "description_eng"
        primary  = (p.get(desc_key) or p.get("description") or "")
        label    = primary[:55] + ("…" if len(primary) > 55 else "")
        PRODUCT_MAP[len(PRODUCT_NAMES)] = p
        PRODUCT_NAMES.append(label)

# ─────────────────────────────────────────────
# DOCX HELPERS
# ─────────────────────────────────────────────
def replace_in_paragraph(para, replacements):
    for key, val in replacements.items():
        full_text = "".join(run.text for run in para.runs)
        if key not in full_text:
            continue
        keeper_run = None
        for run in para.runs:
            if key in run.text or (run.text and run.text in key):
                keeper_run = run
                break
        if keeper_run is None:
            for run in para.runs:
                if run.bold:
                    keeper_run = run
                    break
        if keeper_run is None and para.runs:
            keeper_run = para.runs[-1]
        new_text = full_text.replace(key, val)
        if para.runs:
            para.runs[0].text = new_text
            if keeper_run and keeper_run != para.runs[0]:
                para.runs[0].bold = keeper_run.bold
                para.runs[0].italic = keeper_run.italic
                if keeper_run.font.name:
                    para.runs[0].font.name = keeper_run.font.name
                if keeper_run.font.size:
                    para.runs[0].font.size = keeper_run.font.size
            for run in para.runs[1:]:
                run.text = ""

def set_para_run(para, text, bold=False, font_name="Verdana", font_size=10):
    para.clear()
    r = para.add_run(text)
    r.bold = bold
    r.font.name = font_name
    r.font.size = Pt(font_size)

def set_cell_text(cell, text, bold=False, italic=False, font_name="Verdana", font_size=10):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
            rPr = run._r.find(qn('w:rPr'))
            if rPr is not None:
                run._r.remove(rPr)
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font_name
    run.font.size = Pt(font_size)

# ─────────────────────────────────────────────
# SESSION STATE
# ─────────────────────────────────────────────
if "line_items" not in st.session_state:
    st.session_state.line_items = [
        {"product_idx": 0, "description": "", "details": "", "qty": 1.0,
         "unit_price": 0.0, "price_type": LBL["cliente"]}
    ]

def add_line():
    st.session_state.line_items.append(
        {"product_idx": 0, "description": "", "details": "", "qty": 1.0,
         "unit_price": 0.0, "price_type": LBL["cliente"]}
    )

# ─────────────────────────────────────────────
# UI
# ─────────────────────────────────────────────
col_title, col_lang = st.columns([5, 1])
with col_title:
    st.title(TITLE)
with col_lang:
    st.write("")
    if st.button(LBL["lang_switch"]):
        st.session_state.language = None
        st.session_state.line_items = []
        st.rerun()

# ── 1. DATE & NUMBER ──────────────────────────
st.subheader(f"1. {LBL['date']} & {LBL['number_label']}")
col_d1, col_d2 = st.columns(2)
with col_d1:
    selected_date = st.date_input(LBL["date"], value=date.today(), format="DD/MM/YYYY")
with col_d2:
    suggested_number = get_next_offerta_number()
    year_2digit = selected_date.strftime('%y')
    existing_numbers = load_existing_offerta_numbers()
    proforma_number = st.text_input(
        LBL["number_label"],
        value=suggested_number,
        help=LBL["number_hint"]
    )
    number_ok = True
    if proforma_number in existing_numbers:
        st.error(LBL["number_dup"])
        number_ok = False
    else:
        try:
            seq = int(proforma_number.replace("OF", "").split("/")[0])
            expected = int(suggested_number.replace("OF", "").split("/")[0])
            if seq != expected:
                st.warning(LBL["number_warn"])
        except:
            pass

formatted_date = selected_date.strftime('%d/%m/') + "\u2019" + year_2digit

# ── 2. CLIENT ─────────────────────────────────
st.subheader(LBL["client"])
customers      = st.session_state.customers_db
customer_names = [LBL["new_cust"]] + [
    f"{c.get('company_name', '')} ({c.get('contact_name', '')})" for c in customers
]
col_cust, col_refresh = st.columns([5, 1])
with col_cust:
    selected_customer_idx = st.selectbox(
        LBL["pick_cust"], range(len(customer_names)),
        format_func=lambda x: customer_names[x], key="customer_picker"
    )
with col_refresh:
    st.write("")
    if st.button(LBL["reload"], help="Reload"):
        load_customers.clear()
        st.session_state.customers_db = load_customers()
        st.rerun()

if selected_customer_idx > 0:
    cust = customers[selected_customer_idx - 1]
    sal  = cust.get("salutation", "Mr.") or "Mr."
    default_salutation = sal if sal in ["Mr.", "Ms.", "Dr.", "Messrs."] else "Mr."
    default_full_name  = cust.get("contact_name", "")
    default_company    = cust.get("company_name", "")
    default_address    = cust.get("address", "")
    default_zip        = cust.get("zip", "")
    default_city       = cust.get("city", "")
    default_region     = ""
    default_country    = cust.get("country", "")
else:
    default_salutation = "Mr."
    default_full_name = default_company = default_address = ""
    default_zip = default_city = default_region = default_country = ""

include_attn = st.checkbox(LBL["attn_toggle"], value=True)

if include_attn:
    col1, col2 = st.columns([1, 3])
    with col1:
        salutation = st.selectbox(LBL["salutation"], ["Mr.", "Ms.", "Dr.", "Messrs."],
                                  index=["Mr.", "Ms.", "Dr.", "Messrs."].index(default_salutation))
    with col2:
        full_name = st.text_input(LBL["contact"], value=default_full_name,
                                  placeholder="e.g. John Smith")
else:
    salutation = ""
    full_name  = ""

company = st.text_input(LBL["company"], value=default_company)
address = st.text_input(LBL["address"], value=default_address)
col3, col4, col5 = st.columns(3)
with col3:
    zip_code = st.text_input(LBL["zip"], value=default_zip)
with col4:
    city = st.text_input(LBL["city"], value=default_city)
with col5:
    region = st.text_input(LBL["region"], value=default_region)
country = st.text_input(LBL["country"], value=default_country)

# ── 3. CURRENCY & PRICE TYPE ──────────────────
st.subheader(LBL["currency"])
col_cur, col_pt = st.columns(2)
with col_cur:
    currency_choice = st.selectbox(LBL["cur_lbl"], CURRENCIES)
    if currency_choice == LBL["custom"]:
        currency = st.text_input("ISO code", placeholder="e.g. AED")
    else:
        currency = currency_choice
with col_pt:
    global_price_type = st.radio(
        LBL["price_type"], [LBL["cliente"], LBL["rivenditore"]],
        horizontal=True, key="global_price_type"
    )
    if st.session_state.get("_last_price_type") != global_price_type:
        st.session_state["_last_price_type"] = global_price_type
        for item in st.session_state.line_items:
            item["price_type"] = global_price_type
            if item.get("product_idx", 0) > 0 and item.get("product_idx") in PRODUCT_MAP:
                pc = item.get("price_client", 0.0)
                pr = item.get("price_reseller", 0.0)
                item["unit_price"] = pc if global_price_type == LBL["cliente"] else pr
        st.rerun()

# ── 4. LINE ITEMS ─────────────────────────────
st.subheader(LBL["lines"])
st.caption(LBL["lines_cap"])

items_to_remove = []
needs_rerun = False
for i, item in enumerate(st.session_state.line_items):
    with st.container():
        c1, c2, c3, c4 = st.columns([3, 1.5, 1.5, 0.4])
        with c1:
            prod_idx = st.selectbox(
                LBL["prod"].replace("{i}", str(i+1)),
                range(len(PRODUCT_NAMES)),
                format_func=lambda x: PRODUCT_NAMES[x],
                key=f"prod_{i}", index=item["product_idx"]
            )
            if prod_idx != item["product_idx"]:
                item["product_idx"] = prod_idx
                if prod_idx > 0 and prod_idx in PRODUCT_MAP:
                    p = PRODUCT_MAP[prod_idx]
                    desc_key = "description" if LANG == "it" else "description_eng"
                    item["description"]    = p.get(desc_key) or p.get("description") or ""
                    item["price_client"]   = float(p.get("unit_price_client")   or 0)
                    item["price_reseller"] = float(p.get("unit_price_reseller") or 0)
                    item["unit_price"]     = item["price_client"] if global_price_type == LBL["cliente"] else item["price_reseller"]
                    item["price_type"]     = global_price_type
                else:
                    item["description"] = ""
                    item["unit_price"] = item["price_client"] = item["price_reseller"] = 0.0
                needs_rerun = True

            if prod_idx > 0 and prod_idx in PRODUCT_MAP:
                p_sel = PRODUCT_MAP[prod_idx]
                ita = p_sel.get("description", "")
                eng = p_sel.get("description_eng", "")
                if ita: st.caption(f"🇮🇹 {ita}")
                if eng: st.caption(f"🇬🇧 {eng}")

            item["details"] = st.text_input(
                LBL["details"], value=item.get("details", ""), key=f"details_{i}")

        with c2:
            item["qty"] = st.number_input(
                LBL["qty"], min_value=0.0, value=float(item["qty"]),
                step=1.0, format="%.1f", key=f"qty_{i}")
        with c3:
            st.write(f"**{LBL['unit_price'].format(cur=currency)}**")
            st.write(fmt_price_it(item["unit_price"]))
        with c4:
            st.write("")
            st.write("")
            if st.button(LBL["remove"], key=f"del_{i}"):
                items_to_remove.append(i)

        line_total = item["qty"] * item["unit_price"]
        st.caption(f"Line total: {currency} {fmt_price_it(line_total)}")
        st.divider()

for i in sorted(items_to_remove, reverse=True):
    st.session_state.line_items.pop(i)
if items_to_remove or needs_rerun:
    st.rerun()

st.button(LBL["add_line"], on_click=add_line)
grand_total = sum(item["qty"] * item["unit_price"] for item in st.session_state.line_items)
st.markdown(f"### 💰 Total: {currency} {fmt_price_it(grand_total)}")

# ── 5. TERMS ──────────────────────────────────
st.subheader(LBL["terms"])
DELIVERY_TERMS_OPTIONS = st.session_state.delivery_terms_db
col_t1, col_t2 = st.columns(2)
with col_t1:
    hs_code = st.selectbox(LBL["hs"], HS_CODES + [LBL["custom"]])
    if hs_code == LBL["custom"]:
        hs_code = st.text_input("Custom HS Code")
    payment = st.selectbox(LBL["payment"], PAYMENT_OPTIONS + [LBL["custom"]])
    if payment == LBL["custom"]:
        payment = st.text_input("Custom payment")
    delivery_terms = st.selectbox(LBL["del_terms"], DELIVERY_TERMS_OPTIONS + [LBL["custom"]])
    if delivery_terms == LBL["custom"]:
        delivery_terms = st.text_input("Custom delivery terms", placeholder="e.g. DAP Tokyo")
        if delivery_terms and delivery_terms not in DELIVERY_TERMS_OPTIONS:
            if st.button(LBL["save_dt"], key="save_dt"):
                save_delivery_term(delivery_terms)
                st.session_state.delivery_terms_db = load_delivery_terms()
                st.success(f"✅ '{delivery_terms}' saved!")
                st.rerun()
    delivery_time = st.selectbox(LBL["del_time"], DELIVERY_TIME_OPTIONS + [LBL["custom"]])
    if delivery_time == LBL["custom"]:
        delivery_time = st.text_input("Custom delivery time")
with col_t2:
    packing = st.selectbox(LBL["packing"], PACKING_OPTIONS + [LBL["custom"]])
    if packing == LBL["custom"]:
        packing = st.text_input("Custom packing")
    shipment = st.selectbox(LBL["shipment"], SHIPMENT_OPTIONS + [LBL["custom"]])
    if shipment == LBL["custom"]:
        shipment = st.text_input("Custom shipment")

# ── 6. DOC NAME ───────────────────────────────
st.subheader(LBL["doc_name"])
default_name = f"offerta {proforma_number.replace('/', '-')} {company}"
doc_name = st.text_input(LBL["file_name"], value=default_name)

# ── GENERATE ──────────────────────────────────
st.divider()
if st.button(LBL["generate"], type="primary", use_container_width=True, disabled=not number_ok):
    if not company:
        st.warning(LBL["warn_company"])
    elif not any(item["description"].strip() for item in st.session_state.line_items):
        st.warning(LBL["warn_items"])
    else:
        zip_city = f"{zip_code} {city}".strip()
        if region:
            zip_city += f", {region}"

        header_replacements = {
            f"Schio, [DD/MM/\u2019YY]": f"Schio, {formatted_date}",
            f"[DD/MM/\u2019YY]":        formatted_date,
            "[COMPANY NAME]":           company,
            "[Address]":                address,
            "[Zip] [City], [Region]":   zip_city,
            "[Country]":                country,
            "Mr./Ms. [Full Name]":      f"{salutation} {full_name}" if include_attn else "",
            "[Sal.]":                   salutation,
            "[Full Name]":              full_name,
            "[NNN/YY]":                 proforma_number,
        }

        try:
            template_path = os.path.join(os.path.dirname(__file__), TEMPLATE_FILE)
            doc = Document(template_path)
        except Exception as e:
            st.error(f"❌ Template not found: {e}")
            st.stop()

        for para in doc.paragraphs:
            replace_in_paragraph(para, header_replacements)

        for para in doc.paragraphs:
            full = "".join(r.text for r in para.runs)

            if para == doc.paragraphs[0]:
                para.clear()
                r1 = para.add_run("Schio, ")
                r1.bold = False; r1.font.name = "Verdana"; r1.font.size = Pt(10)
                r2 = para.add_run(formatted_date)
                r2.bold = False; r2.font.name = "Verdana"; r2.font.size = Pt(10)
                continue

            if "To the attn. of" in full or "All'attenzione" in full:
                if include_attn and (salutation or full_name):
                    para.clear()
                    r_prefix = para.add_run(f"To the attn. of {salutation} ")
                    r_prefix.bold = False; r_prefix.font.name = "Verdana"; r_prefix.font.size = Pt(10)
                    r_name = para.add_run(full_name)
                    r_name.bold = False; r_name.font.name = "Verdana"; r_name.font.size = Pt(10)
                else:
                    p = para._p
                    p.getparent().remove(p)
                continue

            if "OFFER NO" in full or "OFFERTA Nr" in full:
                for run in para.runs:
                    run.bold = True
                continue

            if company and company in full:
                set_para_run(para, company, bold=True)
                continue

            for run in para.runs:
                run.bold = False
                run.font.name = "Verdana"
                run.font.size = Pt(10)

        table    = doc.tables[0]
        MAX_ROWS = 15
        valid_items = [it for it in st.session_state.line_items if it["description"].strip()]

        for row_idx in range(1, MAX_ROWS + 1):
            row   = table.rows[row_idx]
            cells = row.cells
            if row_idx - 1 < len(valid_items):
                item       = valid_items[row_idx - 1]
                pos        = row_idx * 10
                line_total = item["qty"] * item["unit_price"]
                qty_str    = f"{item['qty']:,.1f}".replace(",", "X").replace(".", ",").replace("X", ".")
                price_str  = fmt_price_it(item["unit_price"])
                total_str  = fmt_price_it(line_total)
                set_cell_text(cells[0], str(pos), bold=False)
                desc_cell  = cells[1]
                for para in desc_cell.paragraphs:
                    for run in para.runs:
                        run.text = ""
                        rPr = run._r.find(qn('w:rPr'))
                        if rPr is not None: run._r.remove(rPr)
                first_para = desc_cell.paragraphs[0]
                r = first_para.add_run(item["description"])
                r.bold = True; r.font.name = "Verdana"; r.font.size = Pt(10)
                details = item.get("details", "").strip()
                if details:
                    new_p = copy.deepcopy(first_para._p)
                    desc_cell._tc.append(new_p)
                    second_para = desc_cell.paragraphs[-1]
                    for run in second_para.runs:
                        run.text = ""
                    dr = second_para.add_run(details)
                    dr.bold = False; dr.font.name = "Verdana"; dr.font.size = Pt(10)
                set_cell_text(cells[2], qty_str)
                set_cell_text(cells[3], price_str)
                set_cell_text(cells[4], currency)
                set_cell_text(cells[5], total_str)
            else:
                for cell in cells:
                    set_cell_text(cell, "")
                trPr = row._tr.find(qn('w:trPr'))
                if trPr is None:
                    trPr = OxmlElement('w:trPr')
                    row._tr.insert(0, trPr)
                existing_h = trPr.find(qn('w:trHeight'))
                if existing_h is not None:
                    trPr.remove(existing_h)
                trH = OxmlElement('w:trHeight')
                trH.set(qn('w:val'), '1'); trH.set(qn('w:hRule'), 'exact')
                trPr.append(trH)

        total_row   = table.rows[MAX_ROWS + 1]
        tcells      = total_row.cells
        total_label = TOTAL_LABEL_TPL.format(dt=delivery_terms)
        set_cell_text(tcells[0], total_label, bold=True)
        set_cell_text(tcells[4], currency,    bold=True)
        set_cell_text(tcells[5], fmt_price_it(grand_total), bold=True)

        terms_table = doc.tables[1]
        terms_map   = {0: hs_code, 1: payment, 4: delivery_terms,
                       5: delivery_time, 6: packing, 7: shipment}
        for row_idx, value in terms_map.items():
            if row_idx < len(terms_table.rows):
                set_cell_text(terms_table.rows[row_idx].cells[1], value)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        save_offerta(proforma_number, company, grand_total, currency,
                     date_of_reference=selected_date.strftime("%Y-%m-%d"))
        if company.strip():
            save_customer(company, full_name, salutation, "", "", address, city, zip_code, country, "")
            load_customers.clear()
            st.session_state.customers_db = load_customers()

        total_display = fmt_price_it(grand_total)
        st.success(LBL["success"].format(num=proforma_number, cur=currency, total=total_display))
        st.download_button(
            label=LBL["download"], data=buffer,
            file_name=f"{doc_name}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
